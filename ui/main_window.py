import os
import re
import subprocess
import sys
from pathlib import Path

from PyQt6 import sip
from PyQt6.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QGridLayout,
    QPushButton,
    QLabel,
    QFileDialog,
    QLineEdit,
    QComboBox,
    QProgressBar,
    QPlainTextEdit,
    QTabWidget,
    QFormLayout,
    QSpinBox,
    QDoubleSpinBox,
    QMessageBox,
    QInputDialog,
    QScrollArea,
    QSizePolicy,
    QStatusBar,
    QCheckBox,
    QSlider,
    QSplitter,
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtGui import QPixmap, QCursor, QKeySequence, QShortcut, QTextCharFormat, QColor, QSyntaxHighlighter, QPainter, QPen, QTextCursor, QFont
from PyQt6.QtCore import Qt, pyqtSignal, QThread, QObject, QRect, QUrl, QCoreApplication
from ui.pipeline_controller import PipelineController
from ui.comfy_controller import ComfyController
from ui.widgets.workflow_selector import WorkflowSelector
from ui.widgets.node_editor_stub import NodeEditorStub
from ui.widgets.pipeline_status import PipelineStatus
from utilis.config_loader import ConfigLoader


class _DualColorProgressBar(QProgressBar):
    """Progress bar whose percentage label switches from white to black at 50 %."""

    def paintEvent(self, event):
        super().paintEvent(event)
        if not self.isTextVisible():
            return
        text = self.text()
        if not text:
            return
        rect = self.rect()
        fill_ratio = (self.value() - self.minimum()) / max(self.maximum() - self.minimum(), 1)
        fill_px = int(rect.width() * fill_ratio)

        painter = QPainter(self)
        painter.setFont(self.font())

        # Left half (over the filled chunk) — black when >50 %, white when <=50 %
        if fill_ratio > 0.5:
            painter.setPen(QPen(Qt.GlobalColor.black))
        else:
            painter.setPen(QPen(QColor("#E6E1E5")))
        painter.drawText(rect, Qt.AlignmentFlag.AlignCenter, text)
        painter.end()


class _ClickableImageLabel(QLabel):
    clicked = pyqtSignal(str)  # emits the image file path

    def __init__(self, image_path: str, parent=None):
        super().__init__(parent)
        self._image_path = image_path
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setToolTip("Click to enlarge")

    def mousePressEvent(self, event):
        try:
            if event.button() == Qt.MouseButton.LeftButton:
                self.clicked.emit(self._image_path)
            super().mousePressEvent(event)
        except RuntimeError:
            # The label can be deleted mid-click if a background pipeline finishes
            # and rebuilds the grid while this click's dialog is still open/modal.
            pass


class _SpellHighlighter(QSyntaxHighlighter):
    """Underlines misspelled words in red using pyspellchecker."""

    _ITALIAN_ELISIONS = {
        "all", "bell", "c", "coll", "d", "dall", "dell", "gl", "l",
        "m", "n", "nell", "quest", "quell", "s", "senz", "sott", "sull", "t",
        "un",
    }

    # Path to the custom-words file (resolved relative to this source file's
    # grandparent directory, i.e. the repo root → config/).
    _CUSTOM_WORDS_PATH = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
        "config", "spell_custom_words.txt",
    )
    _custom_words: set[str] = set()   # loaded once, shared by all instances

    @classmethod
    def _load_custom_words(cls) -> set[str]:
        """Read the custom word list and remove duplicate entries in place."""
        words: set[str] = set()
        try:
            path = Path(cls._CUSTOM_WORDS_PATH)
            lines = path.read_text(encoding="utf-8-sig").splitlines()
            cleaned_lines: list[str] = []
            for line in lines:
                word = line.split("#", 1)[0].strip()
                normalized = word.casefold()
                if word and normalized in words:
                    continue
                cleaned_lines.append(line)
                if word:
                    words.add(normalized)
            if cleaned_lines != lines:
                path.write_text("\n".join(cleaned_lines).rstrip() + "\n", encoding="utf-8")
        except FileNotFoundError:
            pass
        except Exception:
            pass
        return words

    @classmethod
    def add_custom_word(cls, word: str) -> bool:
        """Append a word to the custom dictionary; return False if it already exists."""
        normalized = word.strip().casefold()
        cls._custom_words = cls._load_custom_words()
        if not normalized or normalized in cls._custom_words:
            return False
        path = Path(cls._CUSTOM_WORDS_PATH)
        path.parent.mkdir(parents=True, exist_ok=True)
        needs_newline = path.exists() and path.stat().st_size > 0
        with path.open("a", encoding="utf-8") as fh:
            if needs_newline:
                with path.open("rb") as source:
                    source.seek(-1, os.SEEK_END)
                    if source.read(1) not in (b"\n", b"\r"):
                        fh.write("\n")
            fh.write(word.strip() + "\n")
        cls._custom_words.add(normalized)
        return True

    def __init__(self, parent, language: str = "it"):
        super().__init__(parent)
        self._fmt = QTextCharFormat()
        self._fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.SpellCheckUnderline)
        self._fmt.setUnderlineColor(QColor("#FF5555"))
        self._checker = None
        # Load custom words once for the class
        if not _SpellHighlighter._custom_words:
            _SpellHighlighter._custom_words = _SpellHighlighter._load_custom_words()
        self.set_language(language)

    def set_language(self, language: str) -> None:
        try:
            from spellchecker import SpellChecker
            self._checker = SpellChecker(language=language)
            if _SpellHighlighter._custom_words:
                self._checker.word_frequency.load_words(_SpellHighlighter._custom_words)
        except Exception:
            self._checker = None
        self.rehighlight()

    def highlightBlock(self, text: str) -> None:
        if not self._checker or not text.strip():
            return
        import re
        for m in re.finditer(r"[A-Za-zÀ-ÖØ-öø-ÿ]+(?:['’][A-Za-zÀ-ÖØ-öø-ÿ]+)?", text):
            word = m.group()
            parts = re.split(r"['’]", word, maxsplit=1)
            checked_word = word
            offset = 0
            if len(parts) == 2 and parts[0].casefold() in self._ITALIAN_ELISIONS:
                checked_word = parts[1]
                offset = len(parts[0]) + 1
            if (checked_word.casefold() not in self._custom_words
                    and self._checker.unknown([checked_word])):
                self.setFormat(m.start() + offset, len(checked_word), self._fmt)


def _load_thumbnail(path: str, w: int, h: int) -> "QPixmap":
    """Load an image from *path* scaled to at most w×h without reading the full
    resolution into memory first.  Uses QImageReader's built-in scaled-read so
    only the required pixel data is decoded — much faster than QPixmap(path).scaled(…)
    for large source images."""
    from PyQt6.QtGui import QImageReader
    reader = QImageReader(path)
    reader.setAutoTransform(True)
    original = reader.size()
    if original.isValid() and not original.isEmpty():
        scaled = original.scaled(
            w, h,
            Qt.AspectRatioMode.KeepAspectRatio,
        )
        reader.setScaledSize(scaled)
    image = reader.read()
    return QPixmap.fromImage(image) if not image.isNull() else QPixmap()


class MainWindow(QMainWindow):
    # Lightbox variant display/generation order (Schnell → Dev → FLUX.2, each −1/●/+1)
    _LIGHTBOX_VARIANT_ORDER = [
        ("schnell_v1", "Schnell −1"),
        ("schnell_v2", "Schnell ●"),
        ("schnell_v3", "Schnell +1"),
        ("dev_v1",     "Dev −1"),
        ("dev_v2",     "Dev ●"),
        ("dev_v3",     "Dev +1"),
        ("flux2_v1",   "FLUX.2 -1"),
        ("flux2_v2",   "FLUX.2 base"),
        ("flux2_v3",   "FLUX.2 +1"),
    ]

    def __init__(self):
        super().__init__()

        self.setWindowTitle("AI Cinematic Video Pipeline")
        self.resize(1100, 680)
        self.setMinimumWidth(1000)
        self.controller = PipelineController()

        central = QWidget()
        layout = QVBoxLayout(central)

        self.tabs = QTabWidget()
        self.tabs.addTab(self._build_pipeline_tab(), "Pipeline")
        self.tabs.addTab(self._build_script_tab(), "Script")
        self.tabs.addTab(self._build_dubbing_tab(), "Dubbing")
        self.tabs.addTab(self._build_prompts_tab(), "Prompts")
        self.tabs.addTab(self._build_draft_tab(), "Preview Images")
        self.tabs.addTab(self._build_lightbox_tab(), "Lightbox")
        self.tabs.addTab(self._build_comfy_tab(), "ComfyUI")
        self.tabs.addTab(self._build_settings_tab(), "Settings")
        layout.addWidget(self.tabs)

        self.setCentralWidget(central)

        # Status bar with Relaunch button
        status_bar = QStatusBar()
        self.setStatusBar(status_bar)
        btn_relaunch = QPushButton("⟳ Relaunch")
        btn_relaunch.setToolTip("Restart the application")
        btn_relaunch.clicked.connect(self._relaunch)
        status_bar.addPermanentWidget(btn_relaunch)

        self._apply_theme()
        self._wire_signals()

        def _ctrl_s():
            tab_title = self.tabs.tabText(self.tabs.currentIndex())
            if tab_title == "Dubbing":
                self._dub_save()
            else:
                self._save_script()
        QShortcut(QKeySequence("Ctrl+S"), self).activated.connect(_ctrl_s)

        self._refresh_recent_projects(self.controller.get_recent_projects())
        self._load_settings_to_form(self.controller.load_settings())

        recent = self.controller.get_recent_projects()
        if recent:
            self.controller.set_project_path(recent[0])

    def _build_pipeline_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)

        project_row = QHBoxLayout()
        self.project_path_input = QLineEdit()
        self.project_path_input.setPlaceholderText("Select or create a project folder")
        btn_select_project = QPushButton("Browse")
        btn_select_project.clicked.connect(self.select_project)
        btn_create_project = QPushButton("Create New Project")
        btn_create_project.clicked.connect(self.create_project)
        project_row.addWidget(QLabel("Project:"))
        project_row.addWidget(self.project_path_input, 1)
        project_row.addWidget(btn_select_project)
        project_row.addWidget(btn_create_project)

        recent_row = QHBoxLayout()
        self.recent_projects_combo = QComboBox()
        self.recent_projects_combo.setPlaceholderText("Recent projects")
        btn_use_recent = QPushButton("Use Selected")
        btn_use_recent.clicked.connect(self.use_recent_project)
        recent_row.addWidget(QLabel("Recent:"))
        recent_row.addWidget(self.recent_projects_combo, 1)
        recent_row.addWidget(btn_use_recent)

        self.pipeline_status_label = QLabel("Ready")
        self.pipeline_progress = _DualColorProgressBar()
        self.pipeline_progress.setRange(0, 100)
        self.pipeline_progress.setValue(0)

        # Individual stage buttons
        stages_row1 = QHBoxLayout()
        stages_row2 = QHBoxLayout()
        self._stage_btns: dict = {}
        stage_defs = [
            ("Preview Images",  "preview_images"),
            ("Preview Clips",   "preview_clips"),
            ("Preview Video",   "preview_video"),
            ("Lightbox Images", "final_images"),
            ("Final Clips",     "final_clips"),
            ("Final Video",     "final_video"),
        ]
        for i, (label, value) in enumerate(stage_defs):
            btn = QPushButton(label)
            btn.clicked.connect(lambda _checked, v=value: self.run_stage(v))
            btn.setMinimumWidth(110)
            self._stage_btns[value] = btn
            if i < 5:
                stages_row1.addWidget(btn)
            else:
                stages_row2.addWidget(btn)

        self.btn_cancel_pipeline = QPushButton("Cancel")
        self.btn_cancel_pipeline.setEnabled(False)
        self.btn_cancel_pipeline.clicked.connect(self.cancel_pipeline)
        stages_row2.addWidget(self.btn_cancel_pipeline)

        extras_row = QHBoxLayout()
        btn_clear_clips = QPushButton("🗑 Clear Final Clips")
        btn_clear_clips.setToolTip("Delete all files in output/clips/")
        btn_clear_clips.clicked.connect(self._clear_clips)
        btn_clear_draft = QPushButton("🗑 Clear Preview Images")
        btn_clear_draft.setToolTip("Delete draft images so they are regenerated")
        btn_clear_draft.clicked.connect(self._clear_draft)
        btn_clear_preview_clips = QPushButton("🗑 Clear Preview Clips")
        btn_clear_preview_clips.setToolTip("Delete all files in output/draft_clips/")
        btn_clear_preview_clips.clicked.connect(self._clear_preview_clips)
        btn_clear_lightbox = QPushButton("🗑 Clear Lightbox")
        btn_clear_lightbox.setToolTip("Delete all generated images in output/lightbox/")
        btn_clear_lightbox.clicked.connect(self._clear_lightbox)
        extras_row.addStretch(1)
        extras_row.addWidget(btn_clear_draft)
        extras_row.addWidget(btn_clear_preview_clips)
        extras_row.addWidget(btn_clear_lightbox)
        extras_row.addWidget(btn_clear_clips)

        self.log_output = QPlainTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setPlaceholderText("Pipeline logs will appear here")

        root.addLayout(project_row)
        root.addLayout(recent_row)
        root.addWidget(self.pipeline_status_label)
        root.addWidget(self.pipeline_progress)
        root.addLayout(stages_row1)
        root.addLayout(stages_row2)
        root.addLayout(extras_row)
        root.addWidget(self.log_output, 1)
        return page

    def _build_comfy_tab(self) -> QWidget:
        workflows_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "workflows")
        config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "config")
        comfy_config = ConfigLoader(config_dir).load(
            "comfy.yaml", default={"host": "127.0.0.1", "port": 8188}
        )
        comfy_host = comfy_config.get("host", "127.0.0.1")
        comfy_port = int(comfy_config.get("port", 8188))
        self.comfy_controller = ComfyController(workflows_dir, host=comfy_host, port=comfy_port)

        tab = QWidget()
        outer_layout = QVBoxLayout(tab)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        left_panel = QWidget()
        layout = QVBoxLayout(left_panel)

        self.comfy_workflow_selector = WorkflowSelector(workflows_dir)
        self.comfy_status = PipelineStatus()

        run_row = QHBoxLayout()
        self.comfy_run_button = QPushButton("Run Workflow")
        self.comfy_cancel_button = QPushButton("Cancel")
        self.comfy_check_button = QPushButton("Check Connection")
        run_row.addWidget(self.comfy_run_button)
        run_row.addWidget(self.comfy_cancel_button)
        run_row.addWidget(self.comfy_check_button)

        layout.addWidget(self.comfy_workflow_selector)
        layout.addLayout(run_row)
        layout.addWidget(self.comfy_status, 1)

        # Embedded ComfyUI web interface, so the node graph lives inside the app
        self.comfy_web_view = QWebEngineView()
        self.comfy_web_view.setUrl(QUrl(f"http://{comfy_host}:{comfy_port}"))
        btn_reload_comfy_ui = QPushButton("Reload ComfyUI Page")
        btn_reload_comfy_ui.clicked.connect(
            lambda: self.comfy_web_view.setUrl(QUrl(f"http://{comfy_host}:{comfy_port}"))
        )
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.addWidget(btn_reload_comfy_ui)
        right_layout.addWidget(self.comfy_web_view, 1)

        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)
        outer_layout.addWidget(splitter)

        self.comfy_check_button.clicked.connect(self.comfy_controller.check_connection)
        self.comfy_run_button.clicked.connect(
            lambda: self.comfy_controller.run_workflow(self.comfy_workflow_selector.selected_workflow())
        )
        self.comfy_cancel_button.clicked.connect(self.comfy_controller.cancel)
        self.comfy_controller.connection_changed.connect(self.comfy_status.set_connected)
        self.comfy_controller.progress.connect(self.comfy_status.set_progress)
        self.comfy_controller.finished.connect(
            lambda result: self.comfy_status.append_log(f"Finished: {result.get('images', [])}")
        )
        self.comfy_controller.failed.connect(lambda err: self.comfy_status.append_log(f"Error: {err}"))

        return tab

    def _build_settings_tab(self) -> QWidget:
        outer = QWidget()
        outer_layout = QVBoxLayout(outer)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        page = QWidget()
        root = QVBoxLayout(page)
        scroll.setWidget(page)
        outer_layout.addWidget(scroll, 1)

        form = QFormLayout()

        from prompts.style_presets import STYLE_PRESETS
        self.style_preset_input = QComboBox()
        self.style_preset_input.addItems(list(STYLE_PRESETS.keys()))

        self.aspect_ratio_input = QComboBox()
        self.aspect_ratio_input.addItems(["16:9", "9:16", "1:1", "4:3", "21:9"])

        self.seed_input = QSpinBox()
        self.seed_input.setRange(0, 2147483647)

        self.fps_input = QSpinBox()
        self.fps_input.setRange(1, 120)

        self.scene_split_method_input = QComboBox()
        self.scene_split_method_input.addItems(["paragraph", "sentence", "semantic", "timed"])

        self.min_sentence_length_input = QSpinBox()
        self.min_sentence_length_input.setRange(1, 500)

        self.sdxl_model_input = QLineEdit()
        self.flux_dev_model_input = QLineEdit()
        self.flux_schnell_model_input = QLineEdit()
        self.flux2_model_input = QLineEdit()
        self.svd_model_input = QLineEdit()

        from images.image_generator import MODEL_TYPES
        self.image_model_input = QComboBox()
        for key, label in MODEL_TYPES.items():
            self.image_model_input.addItem(label, key)
        self.image_model_input.setToolTip("Image generation model architecture")

        self.clip_engine_input = QComboBox()
        self.clip_engine_input.addItems(["ken_burns", "svd"])
        self.clip_engine_input.setToolTip(
            "ken_burns: CPU-only pan/zoom effect (fast, no VRAM)\n"
            "svd: Stable Video Diffusion (GPU, slow, more realistic motion)"
        )

        self.ken_burns_motion_input = QComboBox()
        self.ken_burns_motion_input.addItems(["static", "auto"])
        self.ken_burns_motion_input.setToolTip(
            "static: image is held perfectly still (no zoom or pan)\n"
            "auto: random cinematic pan/zoom applied to each clip"
        )

        self.decode_chunk_size_input = QSpinBox()
        self.decode_chunk_size_input.setRange(1, 25)
        self.decode_chunk_size_input.setValue(4)
        self.decode_chunk_size_input.setToolTip(
            "SVD only: how many frames the VAE decodes at once.\n"
            "Lower = less VRAM (use 2 if you hit out-of-memory), higher = faster."
        )

        self.noise_aug_strength_input = QDoubleSpinBox()
        self.noise_aug_strength_input.setRange(0.0, 1.0)
        self.noise_aug_strength_input.setDecimals(3)
        self.noise_aug_strength_input.setSingleStep(0.01)
        self.noise_aug_strength_input.setValue(0.02)
        self.noise_aug_strength_input.setToolTip(
            "SVD only: how far the clip may drift from the still image.\n"
            "Lower (e.g. 0.0-0.02) = fewer hallucinations / steadier; higher = more motion."
        )

        self.guidance_scale_input = QDoubleSpinBox()
        self.guidance_scale_input.setRange(0.0, 30.0)
        self.guidance_scale_input.setDecimals(2)
        self.guidance_scale_input.setSingleStep(0.5)
        self.guidance_scale_input.setVisible(False)  # legacy, kept for compat

        self.num_inference_steps_input = QSpinBox()
        self.num_inference_steps_input.setRange(1, 200)
        self.num_inference_steps_input.setVisible(False)  # legacy

        # FLUX Schnell settings
        self.schnell_steps_input = QSpinBox()
        self.schnell_steps_input.setRange(1, 20)
        self.schnell_steps_input.setValue(4)
        self.schnell_steps_input.setToolTip("Inference steps for FLUX schnell (Preview Images). Recommended: 4")

        self.schnell_guidance_input = QDoubleSpinBox()
        self.schnell_guidance_input.setRange(0.0, 10.0)
        self.schnell_guidance_input.setDecimals(2)
        self.schnell_guidance_input.setSingleStep(0.5)
        self.schnell_guidance_input.setValue(0.0)
        self.schnell_guidance_input.setToolTip("Guidance scale for FLUX schnell. Use 0.0 (distilled model)")

        # FLUX Dev settings
        self.dev_steps_input = QSpinBox()
        self.dev_steps_input.setRange(1, 100)
        self.dev_steps_input.setValue(20)
        self.dev_steps_input.setToolTip("Inference steps for FLUX dev (Final Images). Recommended: 20–30")

        self.dev_guidance_input = QDoubleSpinBox()
        self.dev_guidance_input.setRange(0.0, 30.0)
        self.dev_guidance_input.setDecimals(2)
        self.dev_guidance_input.setSingleStep(0.5)
        self.dev_guidance_input.setValue(3.5)
        self.dev_guidance_input.setToolTip("Guidance scale for FLUX dev. Recommended: 3.5–5.0")

        self.flux2_steps_input = QSpinBox()
        self.flux2_steps_input.setRange(1, 100)
        self.flux2_steps_input.setValue(4)
        self.flux2_steps_input.setToolTip("Inference steps for FLUX.2 final image variants")

        self.flux2_guidance_input = QDoubleSpinBox()
        self.flux2_guidance_input.setRange(0.0, 30.0)
        self.flux2_guidance_input.setDecimals(2)
        self.flux2_guidance_input.setSingleStep(0.5)
        self.flux2_guidance_input.setValue(1.0)
        self.flux2_guidance_input.setToolTip("Guidance scale for FLUX.2 final image variants")

        self.image_resolution_input = QComboBox()
        self.image_resolution_input.addItem("1024 × 576  (16:9 — fits 16 GiB GPU)", (1024, 576))
        self.image_resolution_input.addItem("1344 × 768  (16:9 — requires 24 GiB GPU)", (1344, 768))
        self.image_resolution_input.setToolTip("Output image resolution for FLUX / SDXL")

        self.num_frames_input = QSpinBox()
        self.num_frames_input.setRange(1, 120)
        self.num_frames_input.setToolTip("SVD only: frames generated per clip (SVD-XT native: 25)")

        self.motion_bucket_id_input = QSpinBox()
        self.motion_bucket_id_input.setRange(0, 255)
        self.motion_bucket_id_input.setToolTip("SVD only: motion intensity (0 = still, 127 = default, 255 = max)")

        self.audio_volume_input = QDoubleSpinBox()
        self.audio_volume_input.setRange(0.0, 3.0)
        self.audio_volume_input.setDecimals(2)
        self.audio_volume_input.setSingleStep(0.1)

        self.fade_in_input = QDoubleSpinBox()
        self.fade_in_input.setRange(0.0, 10.0)
        self.fade_in_input.setDecimals(2)
        self.fade_in_input.setSingleStep(0.1)

        self.fade_out_input = QDoubleSpinBox()
        self.fade_out_input.setRange(0.0, 10.0)
        self.fade_out_input.setDecimals(2)
        self.fade_out_input.setSingleStep(0.1)

        form.addRow("Style preset", self.style_preset_input)
        form.addRow("Aspect ratio", self.aspect_ratio_input)
        form.addRow("Seed", self.seed_input)
        form.addRow("FPS", self.fps_input)
        form.addRow("Scene split method", self.scene_split_method_input)
        form.addRow("Min sentence length", self.min_sentence_length_input)
        form.addRow("SD model path", self.sdxl_model_input)
        form.addRow("FLUX schnell model path", self.flux_schnell_model_input)
        form.addRow("FLUX dev model path", self.flux_dev_model_input)
        form.addRow("FLUX.2 model path", self.flux2_model_input)
        form.addRow("SVD model path", self.svd_model_input)
        form.addRow("Clip engine", self.clip_engine_input)
        form.addRow("Ken Burns motion", self.ken_burns_motion_input)
        form.addRow(QLabel("── FLUX Schnell (Preview Images) ──"))
        form.addRow("Schnell inference steps", self.schnell_steps_input)
        form.addRow("Schnell guidance scale", self.schnell_guidance_input)
        form.addRow(QLabel("── FLUX Dev (Final Images) ──"))
        form.addRow("Dev inference steps", self.dev_steps_input)
        form.addRow("Dev guidance scale", self.dev_guidance_input)
        form.addRow(QLabel("── FLUX.2 (Final Images) ──"))
        form.addRow("FLUX.2 inference steps", self.flux2_steps_input)
        form.addRow("FLUX.2 guidance scale", self.flux2_guidance_input)
        form.addRow(QLabel("── SVD Clip Generation ──"))
        form.addRow("Image resolution", self.image_resolution_input)
        form.addRow("SVD video frames", self.num_frames_input)
        form.addRow("SVD motion bucket id", self.motion_bucket_id_input)
        form.addRow("SVD decode chunk size", self.decode_chunk_size_input)
        form.addRow("SVD noise aug strength", self.noise_aug_strength_input)
        form.addRow("Audio volume", self.audio_volume_input)
        form.addRow("Audio fade in", self.fade_in_input)
        form.addRow("Audio fade out", self.fade_out_input)

        form.addRow(QLabel(""))  # spacer
        form.addRow(QLabel("── Narration TTS (Edge TTS) ──"))

        from narration.tts_engine import EDGE_TTS_VOICES
        self.tts_voice_input = QComboBox()
        for label, short_name, _gender in EDGE_TTS_VOICES:
            self.tts_voice_input.addItem(label, short_name)
        self.tts_voice_input.setToolTip("Voice used for narration synthesis")
        self.tts_voice_input.currentIndexChanged.connect(self._on_tts_voice_changed)

        self._tts_preview_btn = QPushButton("▶ Preview")
        self._tts_preview_btn.setToolTip("Play a sample of the selected voice")
        self._tts_preview_btn.clicked.connect(self._preview_tts_voice)
        voice_row = QHBoxLayout()
        voice_row.setContentsMargins(0, 0, 0, 0)
        voice_row.setSpacing(6)
        voice_row.addWidget(self.tts_voice_input, 1)
        voice_row.addWidget(self._tts_preview_btn)
        voice_widget = QWidget()
        voice_widget.setLayout(voice_row)
        voice_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        form.addRow("TTS voice", voice_widget)

        self.tts_rate_input = QSpinBox()
        self.tts_rate_input.setRange(-50, 100)
        self.tts_rate_input.setValue(0)
        self.tts_rate_input.setSuffix("%")
        self.tts_rate_input.setToolTip("Speaking rate offset (+10 = 10% faster, -10 = 10% slower)")
        self.tts_rate_input.valueChanged.connect(self._tts_rate_changed)
        form.addRow("TTS rate", self.tts_rate_input)

        self.tts_pitch_input = QSpinBox()
        self.tts_pitch_input.setRange(-20, 20)
        self.tts_pitch_input.setValue(0)
        self.tts_pitch_input.setSuffix(" Hz")
        self.tts_pitch_input.setToolTip("Pitch offset in Hz")
        form.addRow("TTS pitch", self.tts_pitch_input)

        self.tts_volume_input = QSpinBox()
        self.tts_volume_input.setRange(-50, 50)
        self.tts_volume_input.setValue(0)
        self.tts_volume_input.setSuffix("%")
        self.tts_volume_input.setToolTip("Volume offset")
        form.addRow("TTS volume", self.tts_volume_input)

        form.addRow(QLabel(""))  # spacer
        form.addRow(QLabel("── Prompt Enhancement (Ollama) ──"))

        self.use_ollama_input = QCheckBox("Enable Ollama prompt enhancement")
        form.addRow("", self.use_ollama_input)

        self.ollama_model_input = QLineEdit()
        self.ollama_model_input.setPlaceholderText("e.g. qwen3:8b")
        form.addRow("Ollama model", self.ollama_model_input)

        self.ollama_host_input = QLineEdit()
        self.ollama_host_input.setPlaceholderText("http://localhost:11434")
        form.addRow("Ollama host", self.ollama_host_input)

        buttons = QHBoxLayout()
        btn_reload = QPushButton("Reload Settings")
        btn_reload.clicked.connect(lambda: self._load_settings_to_form(self.controller.load_settings()))
        btn_save = QPushButton("Save Settings")
        btn_save.clicked.connect(self.save_settings)
        buttons.addWidget(btn_reload)
        buttons.addWidget(btn_save)

        root.addLayout(form)
        root.addLayout(buttons)
        return outer

    def _apply_theme(self) -> None:
        self.setStyleSheet("""
            QMainWindow, QWidget {
                background-color: #0F0D13;
                color: #E6E1E5;
                font-family: 'Segoe UI', sans-serif;
                font-size: 12px;
            }
            QMenuBar {
                background-color: #1D1B20;
                color: #E6E1E5;
                border-bottom: 1px solid #211F26;
            }
            QMenuBar::item:selected { background-color: #2A282F; }
            QMenu {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
            }
            QMenu::item:selected { background-color: #96BDE2; color: #0F0D13; }
            QToolBar {
                background-color: #1D1B20;
                border-bottom: 1px solid #211F26;
                padding: 4px 6px;
                spacing: 4px;
            }
            QPushButton {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
                border-radius: 2px;
                padding: 4px 12px;
                min-height: 24px;
                font-weight: bold;
            }
            QPushButton:hover  { background-color: #2A282F; }
            QPushButton:pressed { background-color: #211F26; }
            QPushButton:disabled { color: #8E8B90; border-color: #211F26; }
            QComboBox {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
                border-radius: 2px;
                padding: 2px 6px;
                min-height: 22px;
            }
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView {
                background-color: #1D1B20;
                color: #E6E1E5;
                selection-background-color: #96BDE2;
                selection-color: #0F0D13;
            }
            QLineEdit {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
                border-radius: 2px;
                padding: 2px 6px;
                min-height: 22px;
            }
            QLineEdit:focus { border: 1px solid #96BDE2; }
            QSpinBox, QDoubleSpinBox {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
                border-radius: 2px;
                padding: 2px 6px;
            }
            QSpinBox::up-button, QDoubleSpinBox::up-button,
            QSpinBox::down-button, QDoubleSpinBox::down-button {
                background-color: #2A282F;
                border: 1px solid #36343B;
                width: 16px;
            }
            QSpinBox::up-button:hover, QDoubleSpinBox::up-button:hover,
            QSpinBox::down-button:hover, QDoubleSpinBox::down-button:hover {
                background-color: #96BDE2;
            }
            QSpinBox::up-arrow, QDoubleSpinBox::up-arrow {
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-bottom: 6px solid #E6E1E5;
                width: 0; height: 0;
            }
            QSpinBox::down-arrow, QDoubleSpinBox::down-arrow {
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 6px solid #E6E1E5;
                width: 0; height: 0;
            }
            QPlainTextEdit {
                background-color: #1D1B20;
                color: #E6E1E5;
                border: 1px solid #211F26;
                border-radius: 2px;
                padding: 4px;
                font-family: 'Consolas', monospace;
                font-size: 11px;
            }
            QPlainTextEdit:focus { border: 1px solid #96BDE2; }
            QTabWidget::pane {
                border: 1px solid #211F26;
                background-color: #0F0D13;
            }
            QTabBar::tab {
                background-color: #1D1B20;
                color: #8E8B90;
                border: 1px solid #211F26;
                padding: 6px 14px;
            }
            QTabBar::tab:selected {
                background-color: #0F0D13;
                color: #E6E1E5;
                border-bottom: 2px solid #96BDE2;
            }
            QTabBar::tab:hover { background-color: #2A282F; }
            QLabel { color: #8E8B90; }
            QProgressBar {
                border: 1px solid #211F26;
                border-radius: 2px;
                text-align: center;
                background: #0F0D13;
                color: #E6E1E5;
            }
            QProgressBar::chunk { background: #96BDE2; border-radius: 2px; }
            QScrollBar:vertical {
                background-color: #211F26; width: 22px; border: none;
            }
            QScrollBar::handle:vertical {
                background-color: #7d9ab9; border-radius: 5px; min-height: 20px;
            }
            QScrollBar::handle:vertical:hover { background-color: #96BDE2; }
            QScrollBar::sub-line:vertical, QScrollBar::add-line:vertical {
                border: none; background: none;
            }
        """)

    def _build_draft_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)
        root.setSpacing(6)

        header_row = QHBoxLayout()
        self._draft_status_label = QLabel("Click \"Run Preview Images\" to generate a storyboard with FLUX schnell.")
        self._draft_status_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        header_row.addWidget(self._draft_status_label, 1)

        btn_run_draft = QPushButton("▶ Run Preview Images")
        btn_run_draft.setToolTip("Generate preview images using FLUX schnell (stage 5)")
        btn_run_draft.clicked.connect(lambda: self.run_stage("preview_images"))
        btn_refresh_draft = QPushButton("Refresh")
        btn_refresh_draft.clicked.connect(self._refresh_draft_grid)
        btn_clear_cache = QPushButton("Clear Prompt Cache")
        btn_clear_cache.setToolTip("Delete prompts.yaml so prompts are regenerated on the next run")
        btn_clear_cache.clicked.connect(self._clear_prompt_cache)
        header_row.addWidget(btn_run_draft)
        header_row.addWidget(btn_refresh_draft)
        header_row.addWidget(btn_clear_cache)
        root.addLayout(header_row)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        self._draft_grid_widget = QWidget()
        self._draft_grid_layout = QVBoxLayout(self._draft_grid_widget)
        self._draft_grid_layout.setSpacing(12)
        self._draft_grid_layout.addStretch(1)
        scroll.setWidget(self._draft_grid_widget)
        root.addWidget(scroll, 1)

        return page


    # ── Lightbox tab ──────────────────────────────────────────────────────────

    def _build_lightbox_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)
        root.setSpacing(6)

        header_row = QHBoxLayout()
        self._lightbox_status_label = QLabel(
            'Run "8. Final Images" to generate three seed variants for every visual beat.')
        self._lightbox_status_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        header_row.addWidget(self._lightbox_status_label, 1)

        btn_run = QPushButton("\u25b6 Run Final Images")
        btn_run.setToolTip("Generate three seed variants per model-specific visual beat")
        btn_run.clicked.connect(lambda: self.run_stage("final_images"))
        btn_refresh = QPushButton("Refresh")
        btn_refresh.clicked.connect(self._refresh_lightbox)
        btn_save = QPushButton("\U0001f4be Save Selections")
        btn_save.setToolTip("Save selected images to lightbox_selections.yaml")
        btn_save.clicked.connect(self._save_lightbox_selections)
        btn_select_all = QPushButton("Select All")
        btn_select_all.clicked.connect(lambda: self._lightbox_set_all(True))
        btn_deselect_all = QPushButton("Deselect All")
        btn_deselect_all.clicked.connect(lambda: self._lightbox_set_all(False))
        self._lightbox_unselected_btn = QPushButton("Hide Unselected")
        self._lightbox_unselected_btn.setCheckable(True)
        self._lightbox_unselected_btn.setToolTip("Hide images that are not selected for Final Clips")
        self._lightbox_unselected_btn.toggled.connect(self._lightbox_toggle_unselected)
        for b in (btn_run, btn_refresh, btn_save, btn_select_all, btn_deselect_all,
                  self._lightbox_unselected_btn):
            header_row.addWidget(b)
        root.addLayout(header_row)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._lightbox_grid_widget = QWidget()
        self._lightbox_grid_layout = QVBoxLayout(self._lightbox_grid_widget)
        self._lightbox_grid_layout.setSpacing(14)
        self._lightbox_grid_layout.addStretch(1)
        scroll.setWidget(self._lightbox_grid_widget)
        root.addWidget(scroll, 1)

        self._lightbox_checkboxes: dict = {}
        self._lightbox_cells: dict = {}
        self._lightbox_scene_cards: dict = {}
        return page

    def _refresh_lightbox(self) -> None:
        import yaml as _yaml

        project = self.project_path_input.text().strip()
        lightbox_dir = os.path.join(project, "output", "lightbox") if project else ""

        while self._lightbox_grid_layout.count() > 1:
            item = self._lightbox_grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._lightbox_checkboxes.clear()
        self._lightbox_cells.clear()
        self._lightbox_scene_cards.clear()

        if not lightbox_dir or not os.path.isdir(lightbox_dir):
            self._lightbox_status_label.setText(
                'No lightbox folder found. Run "8. Final Images" first.')
            return

        scene_files: dict = {}
        for fname in sorted(os.listdir(lightbox_dir)):
            if not fname.lower().endswith(".png"):
                continue
            parts = fname.split("_")
            try:
                sid = int(parts[1])
            except (IndexError, ValueError):
                continue
            scene_files.setdefault(sid, {})[fname] = fname

        if not scene_files:
            self._lightbox_status_label.setText(
                'No lightbox images found. Run "8. Final Images" first.')
            return

        saved_selections: dict = {}
        sel_path = os.path.join(project, "output", "lightbox_selections.yaml")
        if os.path.exists(sel_path):
            try:
                data = _yaml.safe_load(Path(sel_path).read_text(encoding="utf-8")) or {}
                saved_selections = {int(k): set(v) for k, v in data.items()}
            except Exception:
                pass

        scenes_yaml = os.path.join(project, "output", "scenes.yaml")
        scene_texts: dict = {}
        if os.path.exists(scenes_yaml):
            try:
                data = _yaml.safe_load(Path(scenes_yaml).read_text(encoding="utf-8"))
                for s in (data or {}).get("scenes", []):
                    scene_texts[s["id"]] = s["text"]
            except Exception:
                pass

        total_images = sum(len(v) for v in scene_files.values())
        self._lightbox_status_label.setText(
            f"{len(scene_files)} scene(s), {total_images} image(s) \u2014 check images to include in Final Clips")

        THUMB_W, THUMB_H = 180, 102

        def _variant_parts(fname: str):
            match = re.match(
                r"^scene_\d+_(schnell|dev|flux2)_b(\d+)_v(\d+)\.png$", fname, re.IGNORECASE)
            if match:
                model_key, beat, variant = match.groups()
                return model_key.lower(), int(beat), int(variant)
            legacy = re.match(
                r"^scene_\d+_(schnell|dev|flux2)_v(\d+)\.png$", fname, re.IGNORECASE)
            if legacy:
                model_key, variant = legacy.groups()
                return model_key.lower(), 1, int(variant)
            return "other", 999, 999

        model_rank = {"schnell": 0, "dev": 1, "flux2": 2, "other": 3}

        def _variant_order(fname: str):
            model_key, beat, variant = _variant_parts(fname)
            return model_rank[model_key], beat, variant, fname

        for sid in sorted(scene_files.keys()):
            selected_set = saved_selections.get(sid, set())

            scene_card = QWidget()
            scene_card.setStyleSheet(
                "background:#1D1B20; border:1px solid #36343B; border-radius:4px;")
            scene_vlay = QVBoxLayout(scene_card)
            scene_vlay.setContentsMargins(10, 8, 10, 8)
            scene_vlay.setSpacing(4)

            hdr = QLabel(f"Scene {sid}  \u2014  {scene_texts.get(sid, '')[:120]}")
            hdr.setStyleSheet(
                "color:#96BDE2; font-weight:bold; font-size:11px; background:transparent; border:none;")
            hdr.setWordWrap(True)
            scene_vlay.addWidget(hdr)

            thumb_grid = QGridLayout()
            thumb_grid.setSpacing(8)
            self._lightbox_checkboxes[sid] = {}
            self._lightbox_cells[sid] = {}
            self._lightbox_scene_cards[sid] = scene_card

            ordered_files = sorted(scene_files[sid], key=_variant_order)
            for variant_index, fname in enumerate(ordered_files):
                model_key, beat_idx, seed_variant = _variant_parts(fname)
                variant_label = (
                    f"{model_key.upper()} | beat {beat_idx} | seed {seed_variant}"
                    if model_key != "other" else fname
                )
                img_path = os.path.join(lightbox_dir, fname)

                cell = QWidget()
                cell.setFixedWidth(THUMB_W)
                cell.setStyleSheet("background:transparent; border:none;")
                cell_lay = QVBoxLayout(cell)
                cell_lay.setContentsMargins(0, 0, 0, 0)
                cell_lay.setSpacing(2)

                img_lbl = _ClickableImageLabel(img_path)
                img_lbl.clicked.connect(self._open_lightbox_viewer)
                img_lbl.setFixedSize(THUMB_W, THUMB_H)
                img_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                if os.path.exists(img_path):
                    px = QPixmap(img_path)
                    if not px.isNull():
                        img_lbl.setPixmap(px.scaled(
                            THUMB_W, THUMB_H,
                            Qt.AspectRatioMode.KeepAspectRatio,
                            Qt.TransformationMode.SmoothTransformation))
                    else:
                        img_lbl.setText("(error)")
                        img_lbl.setStyleSheet("color:#E73A4B; font-size:10px;")
                else:
                    img_lbl.setText("(missing)")
                    img_lbl.setStyleSheet("color:#8E8B90; font-size:10px;")
                cell_lay.addWidget(img_lbl)

                lbl_row = QHBoxLayout()
                lbl_row.setContentsMargins(0, 0, 0, 0)
                vlbl = QLabel(variant_label)
                vlbl.setStyleSheet(
                    "color:#8E8B90; font-size:10px; background:transparent; border:none;")
                chk = QCheckBox()
                chk.setChecked(fname in selected_set or (not selected_set and os.path.exists(img_path)))
                chk.setToolTip(f"Include {fname} in Final Clips")
                chk.setEnabled(os.path.exists(img_path))
                chk.toggled.connect(
                    lambda _checked, s=sid: self._lightbox_selection_changed(s)
                )
                lbl_row.addWidget(vlbl, 1)
                lbl_row.addWidget(chk)
                cell_lay.addLayout(lbl_row)

                self._lightbox_checkboxes[sid][fname] = chk
                self._lightbox_cells[sid][fname] = cell
                thumb_grid.addWidget(cell, variant_index // 3, variant_index % 3)

            scene_vlay.addLayout(thumb_grid)
            self._lightbox_grid_layout.insertWidget(
                self._lightbox_grid_layout.count() - 1, scene_card)

        # Build ordered flat list used by the navigable viewer.
        self._lightbox_image_list = [
            os.path.join(lightbox_dir, fname)
            for sid in sorted(scene_files.keys())
            for fname in sorted(scene_files[sid], key=_variant_order)
        ]
        self._lightbox_apply_unselected_filter()

    def _lightbox_toggle_unselected(self, hide_unselected: bool) -> None:
        self._lightbox_unselected_btn.setText(
            "Show Unselected" if hide_unselected else "Hide Unselected"
        )
        self._lightbox_apply_unselected_filter()

    def _lightbox_selection_changed(self, sid: int) -> None:
        self._save_lightbox_selections(silent=True)
        if self._lightbox_unselected_btn.isChecked():
            self._lightbox_apply_unselected_filter(sid)

    def _lightbox_apply_unselected_filter(self, only_sid: int = None) -> None:
        hide_unselected = self._lightbox_unselected_btn.isChecked()
        scene_ids = [only_sid] if only_sid is not None else self._lightbox_cells.keys()
        for sid in scene_ids:
            cells = self._lightbox_cells.get(sid, {})
            checkboxes = self._lightbox_checkboxes.get(sid, {})
            for fname, cell in cells.items():
                cell.setVisible(not hide_unselected or checkboxes[fname].isChecked())
            card = self._lightbox_scene_cards.get(sid)
            if card is not None:
                card.setVisible(not hide_unselected or any(
                    checkbox.isChecked() for checkbox in checkboxes.values()
                ))

    def _lightbox_set_all(self, checked: bool) -> None:
        for sid_dict in self._lightbox_checkboxes.values():
            for chk in sid_dict.values():
                if chk.isEnabled():
                    chk.blockSignals(True)
                    chk.setChecked(checked)
                    chk.blockSignals(False)
        self._save_lightbox_selections(silent=True)
        self._lightbox_apply_unselected_filter()

    def _save_lightbox_selections(self, silent: bool = False) -> None:
        import yaml as _yaml

        project = self.project_path_input.text().strip()
        if not project:
            if not silent:
                QMessageBox.warning(self, "No project", "Load a project first.")
            return

        if not self._lightbox_checkboxes:
            if not silent:
                QMessageBox.information(self, "Nothing to save",
                                        "Refresh the Lightbox tab first to load images.")
            return

        selections: dict = {}
        def _order_key(fname: str):
            match = re.match(
                r"^scene_\d+_(schnell|dev|flux2)(?:_b(\d+))?_v(\d+)\.png$",
                fname, re.IGNORECASE)
            if not match:
                return 3, 999, 999, fname
            model_key, beat, variant = match.groups()
            return {"schnell": 0, "dev": 1, "flux2": 2}[model_key.lower()], int(beat or 1), int(variant), fname

        for sid, fname_dict in sorted(self._lightbox_checkboxes.items()):
            chosen = [fname for fname in sorted(fname_dict.keys(), key=_order_key)
                      if fname_dict[fname].isChecked()]
            if chosen:
                selections[sid] = chosen

        sel_path = os.path.join(project, "output", "lightbox_selections.yaml")
        with open(sel_path, "w", encoding="utf-8") as fh:
            _yaml.safe_dump(selections, fh, sort_keys=True, allow_unicode=True)

        total = sum(len(v) for v in selections.values())
        self._lightbox_status_label.setText(
            f"Saved {total} selection(s) across {len(selections)} scene(s) \u2192 lightbox_selections.yaml")
        if not silent:
            self._append_log(f"Lightbox selections saved: {total} image(s) in {len(selections)} scene(s).")

    def _delete_draft_image(self, image_path: str, card: QWidget) -> None:
        try:
            os.remove(image_path)
        except Exception as exc:
            QMessageBox.warning(self, "Delete failed", str(exc))
            return
        card.setParent(None)
        card.deleteLater()
        # Refresh grid to recount and update label
        self._refresh_draft_grid()

    def _open_image_viewer(self, image_path: str) -> None:
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QScrollArea
        pixmap = QPixmap(image_path)
        if pixmap.isNull():
            return

        screen = self.screen().availableGeometry()
        max_w = int(screen.width() * 0.85)
        max_h = int(screen.height() * 0.85)
        scaled = pixmap.scaled(max_w, max_h, Qt.AspectRatioMode.KeepAspectRatio,
                               Qt.TransformationMode.SmoothTransformation)

        dlg = QDialog(self)
        dlg.setWindowTitle(os.path.basename(image_path))
        dlg.setModal(True)
        dlg.resize(scaled.width(), scaled.height())
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(0, 0, 0, 0)
        lbl = QLabel()
        lbl.setPixmap(scaled)
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.mousePressEvent = lambda _: dlg.accept()
        lbl.setToolTip("Click to close")
        lay.addWidget(lbl)
        dlg.exec()

    def _open_lightbox_viewer(self, image_path: str) -> None:
        """Navigable lightbox viewer. Left/Right arrow keys browse all lightbox images."""
        from PyQt6.QtWidgets import (
            QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QCheckBox
        )
        from PyQt6.QtCore import Qt as _Qt
        from PyQt6.QtGui import QKeyEvent

        all_images: list = getattr(self, "_lightbox_image_list", [])
        hide_unselected = self._lightbox_unselected_btn.isChecked()
        image_list = [
            path for path in all_images
            if not hide_unselected
            or (
                (checkbox := self._lightbox_checkboxes
                 .get(int(os.path.basename(path).split("_")[1]), {})
                 .get(os.path.basename(path))) is not None
                and checkbox.isChecked()
            )
        ]
        if not image_list:
            self._open_image_viewer(image_path)
            return

        try:
            start_idx = image_list.index(image_path)
        except ValueError:
            start_idx = 0

        screen = self.screen().availableGeometry()
        max_w = int(screen.width()  * 0.88)
        max_h = int(screen.height() * 0.80)

        dlg = QDialog(self)
        dlg.setModal(True)
        dlg.setStyleSheet(
            "QDialog   { background:#0F0D13; }"
            "QLabel    { color:#E6E1E5; background:transparent; border:none; }"
            "QCheckBox { color:#E6E1E5; font-size:13px; spacing:6px; }"
            "QCheckBox::indicator { width:16px; height:16px; }"
        )
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(8, 8, 8, 8)
        lay.setSpacing(6)

        img_lbl = QLabel()
        img_lbl.setAlignment(_Qt.AlignmentFlag.AlignCenter)
        lay.addWidget(img_lbl, 1)

        metadata_lbl = QLabel()
        metadata_lbl.setAlignment(_Qt.AlignmentFlag.AlignCenter)
        metadata_lbl.setStyleSheet("color:#E6E1E5; font-size:13px; font-weight:bold;")
        lay.addWidget(metadata_lbl)

        # ── Selection checkbox row ────────────────────────────────────────────
        chk_row = QHBoxLayout()
        chk_row.addStretch(1)
        sel_chk = QCheckBox("Include in Final Clips")
        sel_chk.setFocusPolicy(_Qt.FocusPolicy.NoFocus)
        chk_row.addWidget(sel_chk)
        chk_row.addStretch(1)
        lay.addLayout(chk_row)

        # ── Navigation row ────────────────────────────────────────────────────
        nav_row = QHBoxLayout()
        btn_prev = QPushButton("◀  Prev")
        btn_next = QPushButton("Next  ▶")
        info_lbl = QLabel()
        info_lbl.setAlignment(_Qt.AlignmentFlag.AlignCenter)
        info_lbl.setStyleSheet("color:#8E8B90; font-size:11px;")
        for btn in (btn_prev, btn_next):
            btn.setStyleSheet(
                "QPushButton { background:#36343B; color:#E6E1E5; border:none;"
                " border-radius:4px; padding:5px 18px; }"
                "QPushButton:hover { background:#4a4850; }"
                "QPushButton:disabled { color:#555; }"
            )
            btn.setFocusPolicy(_Qt.FocusPolicy.NoFocus)
        nav_row.addWidget(btn_prev)
        nav_row.addStretch(1)
        nav_row.addWidget(info_lbl, 2)
        nav_row.addStretch(1)
        nav_row.addWidget(btn_next)
        lay.addLayout(nav_row)

        state = {"idx": start_idx}

        def _grid_checkbox(path: str):
            """Return the grid QCheckBox for this image path, or None."""
            fname = os.path.basename(path)
            try:
                sid = int(fname.split("_")[1])
            except Exception:
                return None
            return (getattr(self, "_lightbox_checkboxes", {})
                    .get(sid, {}).get(fname))

        def _image_metadata(path: str):
            match = re.match(
                r"^scene_(\d+)_(schnell|dev|flux2)(?:_b(\d+))?_v\d+\.png$",
                os.path.basename(path), re.IGNORECASE,
            )
            if not match:
                return None
            scene, model_key, beat = match.groups()
            model_name = {"schnell": "Schnell", "dev": "DEV", "flux2": "FLUX.2"}[
                model_key.lower()
            ]
            return int(scene), int(beat or 1), model_name

        # Sync the viewer checkbox → grid checkbox (one direction)
        def _on_viewer_chk(checked: int):
            grid_chk = _grid_checkbox(image_list[state["idx"]])
            if grid_chk is not None and grid_chk.isEnabled():
                grid_chk.setChecked(bool(checked))

        sel_chk.stateChanged.connect(_on_viewer_chk)

        def _load(idx: int):
            idx = max(0, min(len(image_list) - 1, idx))
            state["idx"] = idx
            path = image_list[idx]
            px = QPixmap(path)
            if not px.isNull():
                scaled = px.scaled(max_w, max_h,
                                   _Qt.AspectRatioMode.KeepAspectRatio,
                                   _Qt.TransformationMode.SmoothTransformation)
                img_lbl.setPixmap(scaled)
                dlg.resize(max(scaled.width() + 16, 480),
                            scaled.height() + 90)
            else:
                img_lbl.setText("(cannot load image)")
            name = os.path.basename(path)
            dlg.setWindowTitle(name)
            metadata = _image_metadata(path)
            metadata_lbl.setText(
                f"Scene {metadata[0]}  |  Beat {metadata[1]}  |  {metadata[2]}"
                if metadata else name
            )
            info_lbl.setText(
                f"{name}   [{idx + 1} / {len(image_list)}]"
                "   ◀ ▶  or  ← →  to navigate   ·  Esc to close"
            )
            btn_prev.setEnabled(idx > 0)
            btn_next.setEnabled(idx < len(image_list) - 1)

            # Mirror grid checkbox state into viewer checkbox (block signal to avoid loop)
            grid_chk = _grid_checkbox(path)
            sel_chk.blockSignals(True)
            if grid_chk is not None:
                sel_chk.setChecked(grid_chk.isChecked())
                sel_chk.setEnabled(grid_chk.isEnabled())
            else:
                sel_chk.setChecked(False)
                sel_chk.setEnabled(False)
            sel_chk.blockSignals(False)

        btn_prev.clicked.connect(lambda: _load(state["idx"] - 1))
        btn_next.clicked.connect(lambda: _load(state["idx"] + 1))

        def _key(event: QKeyEvent):
            key = event.key()
            if key in (_Qt.Key.Key_Right, _Qt.Key.Key_Down):
                _load(state["idx"] + 1)
            elif key in (_Qt.Key.Key_Left, _Qt.Key.Key_Up):
                _load(state["idx"] - 1)
            elif key == _Qt.Key.Key_Escape:
                dlg.accept()
            else:
                QDialog.keyPressEvent(dlg, event)

        dlg.keyPressEvent = _key
        _load(start_idx)
        dlg.exec()

        # Auto-save selections after the viewer closes (X button or Esc) so
        # that any checkbox changes made inside the viewer are persisted to
        # lightbox_selections.yaml immediately — without requiring the user
        # to manually click the Save button afterwards.
        self._save_lightbox_selections()

    def _open_draft_zoom_dialog(self, image_path: str) -> None:
        """Zoom dialog for a draft image: shows zoomed image, narration text (read-only),
        editable prompt and a Save / Redo button that regenerates just that image."""
        import yaml
        from PyQt6.QtWidgets import (
            QDialog, QVBoxLayout, QHBoxLayout, QLabel,
            QTextEdit, QPushButton, QFrame,
        )

        fname = os.path.basename(image_path)
        try:
            scene_id = int(fname.split("_")[1].split(".")[0])
        except Exception:
            scene_id = None

        project = self.project_path_input.text().strip()
        scenes_yaml  = os.path.join(project, "output", "scenes.yaml")
        prompts_yaml = os.path.join(project, "output", "prompts.yaml")
        overrides_yaml = os.path.join(project, "output", "prompt_overrides.yaml")

        # Load narration text
        scene_text = ""
        if scene_id is not None and os.path.exists(scenes_yaml):
            try:
                data = yaml.safe_load(Path(scenes_yaml).read_text(encoding="utf-8"))
                for s in (data or {}).get("scenes", []):
                    if s["id"] == scene_id:
                        scene_text = s.get("text", "")
                        break
            except Exception:
                pass

        # Load prompt
        prompts: dict = {}
        overrides: dict = {}
        current_prompt = ""
        if scene_id is not None and os.path.exists(prompts_yaml):
            try:
                prompts = yaml.safe_load(Path(prompts_yaml).read_text(encoding="utf-8")) or {}
                current_prompt = prompts.get(scene_id) or prompts.get(str(scene_id)) or ""
            except Exception:
                pass
        if scene_id is not None and os.path.exists(overrides_yaml):
            try:
                overrides = yaml.safe_load(Path(overrides_yaml).read_text(encoding="utf-8")) or {}
                current_prompt = (overrides.get(scene_id) or overrides.get(str(scene_id))
                                  or current_prompt)
            except Exception:
                pass

        pixmap = QPixmap(image_path)
        if pixmap.isNull():
            return

        screen = self.screen()
        if screen is None:
            screen = self.app.primaryScreen() if hasattr(self, 'app') else None
        if screen is None:
            from PyQt6.QtWidgets import QApplication
            screen = QApplication.primaryScreen()
        screen_geom = screen.availableGeometry() if screen is not None else None
        if screen_geom is None:
            img_max_w, img_max_h = 800, 450
        else:
            img_max_w = int(screen_geom.width()  * 0.60)
            img_max_h = int(screen_geom.height() * 0.45)
        scaled = pixmap.scaled(img_max_w, img_max_h,
                               Qt.AspectRatioMode.KeepAspectRatio,
                               Qt.TransformationMode.SmoothTransformation)


        dlg = QDialog(self)
        dlg.setWindowTitle(f"Scene {scene_id}  —  {fname}")
        dlg.setModal(True)
        # WA_DeleteOnClose ensures Qt destroys the C++ QDialog object when it
        # closes, rather than letting it accumulate as a hidden child of
        # MainWindow.  Without this, repeated zoom sessions pile up dead
        # QDialog instances which can eventually cause a segfault.
        dlg.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
        dlg.setMinimumWidth(max(scaled.width(), 620))
        dlg.setStyleSheet(
            "QDialog { background:#1D1B20; }"
            "QLabel  { color:#E6E1E5; background:transparent; border:none; }"
        )

        lay = QVBoxLayout(dlg)
        lay.setSpacing(8)
        lay.setContentsMargins(14, 14, 14, 14)

        # ── Zoomed image ─────────────────────────────────────────────────────
        img_lbl = QLabel()
        img_lbl.setPixmap(scaled)
        img_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lay.addWidget(img_lbl)

        def _sep():
            f = QFrame()
            f.setFrameShape(QFrame.Shape.HLine)
            f.setStyleSheet("QFrame { color:#36343B; background:#36343B; max-height:1px; }")
            return f

        lay.addWidget(_sep())

        # ── Narration text (read-only) ────────────────────────────────────────
        narr_hdr = QLabel("Narration")
        narr_hdr.setStyleSheet("color:#96BDE2; font-weight:bold; font-size:11px;")
        lay.addWidget(narr_hdr)

        narr_view = QTextEdit()
        narr_view.setPlainText(scene_text)
        narr_view.setReadOnly(True)
        narr_view.setFixedHeight(70)
        narr_view.setStyleSheet(
            "QTextEdit { background:#0F0D13; color:#E6E1E5; border:1px solid #36343B;"
            " border-radius:3px; font-size:12px; padding:4px; }"
        )
        lay.addWidget(narr_view)

        lay.addWidget(_sep())

        # ── Prompt (editable) ────────────────────────────────────────────────
        prompt_hdr = QLabel("Image Prompt  (edit then click  Save & Redo  to regenerate)")
        prompt_hdr.setStyleSheet("color:#96BDE2; font-weight:bold; font-size:11px;")
        lay.addWidget(prompt_hdr)

        prompt_edit = QTextEdit()
        prompt_edit.setPlainText(current_prompt)
        prompt_edit.setFixedHeight(110)
        prompt_edit.setStyleSheet(
            "QTextEdit { background:#0F0D13; color:#E6E1E5; border:1px solid #36343B;"
            " border-radius:3px; font-size:14px; padding:4px; }"
            "QTextEdit:focus { border:1px solid #96BDE2; }"
        )
        lay.addWidget(prompt_edit)

        # ── Button row ───────────────────────────────────────────────────────
        btn_row = QHBoxLayout()
        btn_row.addStretch(1)

        btn_close = QPushButton("Close")
        btn_close.setStyleSheet(
            "QPushButton { background:#36343B; color:#E6E1E5; border:none;"
            " border-radius:4px; padding:6px 20px; }"
            "QPushButton:hover { background:#4a4850; }"
        )
        btn_close.clicked.connect(dlg.reject)

        btn_save = QPushButton("\U0001f4be  Save & Redo Image")
        btn_save.setToolTip("Save the edited prompt and regenerate this draft image only")
        btn_save.setStyleSheet(
            "QPushButton { background:#2A6099; color:white; font-weight:bold;"
            " border:none; border-radius:4px; padding:6px 20px; }"
            "QPushButton:hover { background:#3578B8; }"
        )

        btn_save_clip = QPushButton("\U0001f3ac  Save & Redo Image + Clip")
        btn_save_clip.setToolTip(
            "Save the edited prompt, regenerate the draft image, then regenerate the preview clip")
        btn_save_clip.setStyleSheet(
            "QPushButton { background:#256040; color:white; font-weight:bold;"
            " border:none; border-radius:4px; padding:6px 20px; }"
            "QPushButton:hover { background:#2e7a50; }"
        )

        btn_reset = QPushButton("\u21bb  Reset to Auto-Generated")
        btn_reset.setToolTip("Discard manual edits and rebuild the prompt from current Settings")
        btn_reset.setStyleSheet(
            "QPushButton { background:#36343B; color:#E6E1E5; border:none;"
            " border-radius:4px; padding:6px 20px; }"
            "QPushButton:hover { background:#4a4850; }"
        )

        def _reset_to_auto_generated():
            if scene_id is None:
                return
            from prompts.prompt_builder import PromptBuilder
            builder = PromptBuilder(
                style_preset=self.style_preset_input.currentText(),
                default_aspect_ratio=self.aspect_ratio_input.currentText(),
                use_ollama=self.use_ollama_input.isChecked(),
                ollama_model=self.ollama_model_input.text().strip() or "qwen3:8b",
                ollama_host=self.ollama_host_input.text().strip() or "http://localhost:11434",
            )
            rebuilt = builder.build_prompt({"id": scene_id, "text": scene_text})
            prompt_edit.setPlainText(rebuilt)

        btn_reset.clicked.connect(_reset_to_auto_generated)

        def _persist_prompt_and_delete_image():
            """Save prompt to yaml and delete the draft image. Returns True on success."""
            new_prompt = prompt_edit.toPlainText().strip()
            if scene_id is None:
                return False
            overrides[scene_id] = new_prompt
            overrides.pop(str(scene_id), None)
            try:
                with open(overrides_yaml, "w", encoding="utf-8") as fh:
                    yaml.safe_dump(overrides, fh, allow_unicode=True, sort_keys=False)
            except Exception as exc:
                QMessageBox.warning(dlg, "Save failed", str(exc))
                return False
            if os.path.exists(image_path):
                try:
                    os.remove(image_path)
                except Exception as exc:
                    QMessageBox.warning(dlg, "Delete failed", str(exc))
                    return False
            return True

        def _save_and_redo():
            if not _persist_prompt_and_delete_image():
                return
            dlg.accept()
            self.run_stage("preview_images")

        def _save_and_redo_clip():
            if not _persist_prompt_and_delete_image():
                return
            # Also delete the matching draft clip so preview_clips regenerates it.
            clip_path = os.path.join(
                self.project_path_input.text().strip(),
                "output", "draft_clips", f"scene_{scene_id:03d}.mp4"
            )
            if os.path.exists(clip_path):
                try:
                    os.remove(clip_path)
                except Exception as exc:
                    QMessageBox.warning(dlg, "Delete clip failed", str(exc))
                    return
            dlg.accept()
            # Chain: run preview_images, then automatically run preview_clips after.
            def _chain(success: bool, payload: str):
                try:
                    self.controller.pipeline_finished.disconnect(_chain)
                except Exception:
                    return  # already disconnected (e.g. button double-clicked)
                if success:
                    self.run_stage("preview_clips")
            self.controller.pipeline_finished.connect(_chain)
            self.run_stage("preview_images")

        btn_save.clicked.connect(_save_and_redo)
        btn_save_clip.clicked.connect(_save_and_redo_clip)
        btn_row.addWidget(btn_reset)
        btn_row.addWidget(btn_close)
        btn_row.addWidget(btn_save)
        btn_row.addWidget(btn_save_clip)
        lay.addLayout(btn_row)

        # Suppress grid rebuilds while this dialog's nested event loop is running —
        # a background pipeline finishing mid-dialog would otherwise delete the very
        # card widget whose click opened this dialog, crashing on return to its
        # mousePressEvent. See _refresh_draft_grid().
        self._draft_dialog_depth = getattr(self, "_draft_dialog_depth", 0) + 1
        try:
            dlg.exec()
        finally:
            self._draft_dialog_depth -= 1
        if self._draft_dialog_depth <= 0:
            self._draft_grid_refresh_pending = False
            self._refresh_draft_grid()

    def _refresh_draft_grid(self) -> None:
        if getattr(self, "_draft_dialog_depth", 0) > 0:
            # A draft/prompt dialog's nested event loop is active — rebuilding now
            # would delete widgets still referenced further up the call stack.
            # Defer until the dialog closes (see _open_draft_zoom_dialog).
            self._draft_grid_refresh_pending = True
            return
        project = self.project_path_input.text().strip()
        draft_dir   = os.path.join(project, "output", "draft")  if project else ""
        images_dir  = os.path.join(project, "output", "images") if project else ""

        # Prefer full-res images when they exist; fall back to draft folder
        def _images_from(folder: str):
            if not folder or not os.path.isdir(folder):
                return []
            return sorted(f for f in os.listdir(folder)
                          if f.lower().endswith((".png", ".jpg", ".jpeg")))

        final_images = _images_from(images_dir)
        draft_images = _images_from(draft_dir)

        # Build merged per-scene lists: prefer final images when present, otherwise
        # retain every draft beat/variant belonging to the scene.
        def _scene_id_from(fname: str) -> int:
            try:
                return int(fname.split("_")[1].split(".")[0])
            except Exception:
                return 0

        scene_entries: dict = {}  # scene_id -> [(dir, fname), ...]
        for fname in draft_images:
            sid = _scene_id_from(fname)
            if sid > 0:
                scene_entries.setdefault(sid, []).append((draft_dir, fname))
        final_scene_ids = {_scene_id_from(fname) for fname in final_images}
        for sid in final_scene_ids:
            if sid > 0:
                scene_entries[sid] = []
        for fname in final_images:
            sid = _scene_id_from(fname)
            if sid > 0:
                scene_entries.setdefault(sid, []).append((images_dir, fname))

        entries = [
            (sid, directory, fname)
            for sid, scene_images in sorted(scene_entries.items())
            for directory, fname in scene_images
        ]
        n_final = sum(1 for _, directory, _ in entries if directory == images_dir)
        n_draft = len(entries) - n_final
        source_label = (
            f"{n_final} final + {n_draft} draft" if n_final and n_draft
            else ("final images" if n_final else "draft images")
        )

        # clear existing cards
        while self._draft_grid_layout.count() > 1:
            item = self._draft_grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not scene_entries:
            self._draft_status_label.setText("No images found. Run Preview Images (stage 5) first.")
            return

        scenes_yaml = os.path.join(project, "output", "scenes.yaml")
        scene_texts: dict = {}
        if os.path.exists(scenes_yaml):
            try:
                import yaml
                data = yaml.safe_load(Path(scenes_yaml).read_text(encoding="utf-8"))
                for s in (data or {}).get("scenes", []):
                    scene_texts[s["id"]] = s["text"]
            except Exception:
                pass

        self._draft_status_label.setText(f"{len(entries)} image(s) — showing {source_label}")

        for sid, active_dir, fname in entries:
            scene_id = sid

            card = QWidget()
            card.setStyleSheet("background:#1D1B20; border:1px solid #36343B; border-radius:4px;")
            card_layout = QHBoxLayout(card)
            card_layout.setContentsMargins(10, 8, 10, 8)
            card_layout.setSpacing(12)

            img_label = _ClickableImageLabel(os.path.join(active_dir, fname))
            img_label.clicked.connect(self._open_draft_zoom_dialog)
            img_label.setFixedSize(160, 90)
            img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            pixmap = _load_thumbnail(os.path.join(active_dir, fname), 160, 90)
            if not pixmap.isNull():
                img_label.setPixmap(pixmap)
            else:
                img_label.setText("(no image)")
            img_label.setToolTip("Click to zoom, edit prompt or regenerate")
            card_layout.addWidget(img_label)

            text_block = QVBoxLayout()
            beat_match = re.search(r"_schnell_b(\d+)_v(\d+)\.", fname, re.IGNORECASE)
            image_detail = (
                f" | Beat {int(beat_match.group(1))} | Seed {int(beat_match.group(2))}"
                if beat_match else ""
            )
            id_lbl = QLabel(f"Scene {scene_id}{image_detail}")
            id_lbl.setStyleSheet("color:#96BDE2; font-weight:bold; font-size:11px; background:transparent; border:none;")
            text_lbl = QLabel(scene_texts.get(scene_id, ""))
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("color:#E6E1E5; font-size:12px; background:transparent; border:none;")
            text_lbl.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
            text_block.addWidget(id_lbl)
            text_block.addWidget(text_lbl)
            text_block.addStretch(1)
            card_layout.addLayout(text_block, 1)

            image_file = os.path.join(active_dir, fname)
            btn_delete = QPushButton("🗑")
            btn_delete.setToolTip("Delete this draft image")
            btn_delete.setFixedSize(28, 28)
            btn_delete.setStyleSheet(
                "QPushButton { color:#E73A4B; background:#1D1B20; border:1px solid #36343B; "
                "border-radius:3px; font-size:14px; padding:0; }"
                "QPushButton:hover { background:#2A282F; }"
            )
            btn_delete.clicked.connect(lambda checked, p=image_file, c=card: self._delete_draft_image(p, c))
            card_layout.addWidget(btn_delete)

            self._draft_grid_layout.insertWidget(self._draft_grid_layout.count() - 1, card)

    # ── Dubbing tab ───────────────────────────────────────────────────────────
    # Color palette (matches dubbing_editor.py)
    _DUB_NOT_DUBBED = ("#0e2a3d", "#5BB4D8")   # dark-blue bg, light-blue icon
    _DUB_TO_REDUB   = ("#3d3000", "#F8B23D")   # amber  — text changed after dub
    _DUB_DUBBED     = ("#0d2e14", "#4CAF50")   # green  — audio up-to-date

    _DUB_BTN_COLORS = {
        "none":      ("#555555", "#444444"),   # grey   — no segments loaded
        "all_stale": ("#c0392b", "#a93226"),   # red    — nothing dubbed yet
        "stale":     ("#e67e22", "#ca6f1e"),   # orange — some need redubbing
        "complete":  ("#27ae60", "#1e8449"),   # green  — all up to date
    }
    _BTN_SS = (
        "QPushButton {{ background-color:{bg}; color:white; font-weight:bold; "
        "font-size:13px; border-radius:5px; padding:4px 14px; border:none; }}"
        "QPushButton:hover {{ background-color:{hv}; }}"
        "QPushButton:pressed {{ background-color:{hv}; }}"
        "QPushButton:disabled {{ background-color:#555; color:#999; }}"
    )

    def _build_dubbing_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)
        root.setSpacing(4)

        # ── Toolbar ──────────────────────────────────────────────────────────
        bar = QHBoxLayout()
        self._dub_status_label = QLabel("Load scenes to start editing dubbed text.")
        self._dub_status_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        bar.addWidget(self._dub_status_label, 1)

        btn_load = QPushButton("Load from Scenes")
        btn_load.setToolTip("Populate cards from scenes.yaml")
        btn_load.clicked.connect(self._dub_load_from_scenes)

        self._dub_save_btn = QPushButton("💾 Save Dubbing")
        self._dub_save_btn.setToolTip("Save all dubbed text to dubbing.yaml  (Ctrl+S)")
        self._dub_save_btn.clicked.connect(self._dub_save)

        self._dub_all_btn = QPushButton("⏩ Dub All")
        self._dub_all_btn.setToolTip("Synthesise audio for all segments")
        self._dub_all_btn.clicked.connect(self._dub_all)

        self._dub_redub_all_btn = QPushButton("↻ Redub All")
        self._dub_redub_all_btn.setToolTip(
            "Force re-synthesize EVERY segment, even if already dubbed "
            "(use after changing speed, pitch or voice for the whole project)"
        )
        self._dub_redub_all_btn.clicked.connect(lambda: self._dub_all(force=True))

        self._dub_play_btn = QPushButton("▶ Play")
        self._dub_play_btn.setToolTip("Play all dubbed audio starting from the current scene")
        self._dub_play_btn.clicked.connect(self._dub_play_from_current)

        speed_label = QLabel("Speed:")
        self._dub_speed_input = QDoubleSpinBox()
        self._dub_speed_input.setRange(0.5, 2.0)
        self._dub_speed_input.setSingleStep(0.01)
        self._dub_speed_input.setDecimals(2)
        self._dub_speed_input.setValue(1.0)
        self._dub_speed_input.setSuffix("x")
        self._dub_speed_input.setFixedHeight(28)
        self._dub_speed_input.setToolTip("Narration speed (1.0 = normal, 0.95 = 5% slower, 1.10 = 10% faster)")
        self._dub_speed_input.valueChanged.connect(self._dub_speed_changed)

        btn_export_word = QPushButton("📄 Export Word")
        btn_export_word.setToolTip("Export all dubbed text to a .docx file in the project output folder")
        btn_export_word.clicked.connect(self._dub_export_word)

        # Spell-check language toggle (IT / EN — no Off button; spell check always on)
        self._dub_spell_lang = "it"
        self._dub_spell_highlighters: dict = {}  # sid → _SpellHighlighter
        btn_spell_it = QPushButton("🇮🇹 IT")
        btn_spell_en = QPushButton("🇬🇧 EN")
        for b in (btn_spell_it, btn_spell_en):
            b.setFixedHeight(28)
        self._dub_spell_btns = {"it": btn_spell_it, "en": btn_spell_en}
        btn_spell_it.clicked.connect(lambda: self._dub_set_spell_lang("it"))
        btn_spell_en.clicked.connect(lambda: self._dub_set_spell_lang("en"))
        self._dub_update_spell_btn_style()

        btn_find = QPushButton("🔍 Find")
        btn_find.setToolTip("Open Find / Replace panel  (Ctrl+F)")
        btn_find.setFixedHeight(28)
        btn_find.clicked.connect(self._dub_toggle_find_panel)

        self._dub_duration_label = QLabel("00:00")
        self._dub_duration_label.setToolTip("Overall length of dubbed scenes")
        self._dub_duration_label.setAlignment(
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
        )
        self._dub_duration_label.setStyleSheet(
            "color:#E6E1E5; font-size:12px; font-weight:bold;"
        )

        bar.addWidget(btn_load)
        bar.addWidget(self._dub_save_btn)
        bar.addWidget(self._dub_all_btn)
        bar.addWidget(self._dub_redub_all_btn)
        bar.addWidget(self._dub_play_btn)
        bar.addWidget(speed_label)
        bar.addWidget(self._dub_speed_input)
        bar.addWidget(btn_export_word)
        bar.addWidget(btn_find)
        bar.addWidget(btn_spell_it)
        bar.addWidget(btn_spell_en)
        bar.addStretch(1)   # pushes status label to expand on the left
        bar.addWidget(self._dub_duration_label)
        root.addLayout(bar)

        # ── Find / Replace panel (hidden by default) ──────────────────────────
        self._dub_find_panel = QWidget()
        self._dub_find_panel.setVisible(False)
        fp = QHBoxLayout(self._dub_find_panel)
        fp.setContentsMargins(0, 2, 0, 2)
        fp.setSpacing(4)

        self._dub_find_input = QLineEdit()
        self._dub_find_input.setPlaceholderText("Find…")
        self._dub_find_input.setFixedHeight(26)
        self._dub_find_input.returnPressed.connect(self._dub_find_next)
        self._dub_find_input.textChanged.connect(lambda: setattr(self, "_dub_find_cursor", None))

        self._dub_replace_input = QLineEdit()
        self._dub_replace_input.setPlaceholderText("Replace with…")
        self._dub_replace_input.setFixedHeight(26)

        self._dub_find_case = QCheckBox("Match case")
        self._dub_find_word = QCheckBox("Whole word")

        btn_fn = QPushButton("▼ Next")
        btn_fn.setFixedHeight(26)
        btn_fn.clicked.connect(self._dub_find_next)
        btn_fp = QPushButton("▲ Prev")
        btn_fp.setFixedHeight(26)
        btn_fp.clicked.connect(self._dub_find_prev)
        btn_repl = QPushButton("Replace")
        btn_repl.setFixedHeight(26)
        btn_repl.clicked.connect(self._dub_replace_one)
        btn_repl_all = QPushButton("Replace All")
        btn_repl_all.setFixedHeight(26)
        btn_repl_all.clicked.connect(self._dub_replace_all)
        btn_close_find = QPushButton("X")
        btn_close_find.setFixedSize(26, 26)
        btn_close_find.setToolTip("Close find/replace panel")
        btn_close_find.setStyleSheet("QPushButton { color: #E6E1E5; font-weight: bold; }")
        btn_close_find.clicked.connect(self._dub_close_find_panel)

        self._dub_find_status = QLabel("")
        self._dub_find_status.setStyleSheet("color:#8E8B90; font-size:11px;")

        fp.addWidget(QLabel("Find:"))
        fp.addWidget(self._dub_find_input, 2)
        fp.addWidget(QLabel("Replace:"))
        fp.addWidget(self._dub_replace_input, 2)
        fp.addWidget(self._dub_find_case)
        fp.addWidget(self._dub_find_word)
        fp.addWidget(btn_fp)
        fp.addWidget(btn_fn)
        fp.addWidget(btn_repl)
        fp.addWidget(btn_repl_all)
        fp.addWidget(self._dub_find_status, 1)
        fp.addWidget(btn_close_find)
        root.addWidget(self._dub_find_panel)

        # Ctrl+F shortcut to open the panel
        _find_sc = QShortcut(QKeySequence("Ctrl+F"), page)
        _find_sc.activated.connect(self._dub_open_find_panel)

        # ── Progress bar (shown only during Dub All) ──────────────────────────
        self._dub_progress = QProgressBar()
        self._dub_progress.setRange(0, 100)
        self._dub_progress.setFixedHeight(6)
        self._dub_progress.setTextVisible(False)
        self._dub_progress.setStyleSheet(
            "QProgressBar { background:#0F0D13; border:none; }"
            "QProgressBar::chunk { background:#4CAF50; }"
        )
        self._dub_progress.setVisible(False)
        root.addWidget(self._dub_progress)

        # ── Segment cards ─────────────────────────────────────────────────────
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._dub_cards_widget = QWidget()
        self._dub_cards_layout = QVBoxLayout(self._dub_cards_widget)
        self._dub_cards_layout.setSpacing(6)
        self._dub_cards_layout.addStretch(1)
        scroll.setWidget(self._dub_cards_widget)
        self._dub_scroll_area = scroll
        root.addWidget(scroll, 1)

        return page

    def _build_prompts_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)
        root.setSpacing(6)

        self._prompts_status_label = QLabel("Build prompts or reload translations to populate scenes.")
        self._prompts_status_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        root.addWidget(self._prompts_status_label)

        profile_bar = QHBoxLayout()
        profile_bar.addWidget(QLabel("Project Aesthetic Profile:"))
        self._project_profile_combo = QComboBox()
        self._project_profile_combo.setToolTip(
            "Geographic, historical, or stylistic constraints appended to every model's "
            "Qwen instructions (applies uniformly to Schnell, Dev, and FLUX.2)."
        )
        self._project_profile_combo.currentIndexChanged.connect(self._on_project_profile_selected)
        profile_bar.addWidget(self._project_profile_combo, 1)
        btn_manage_profiles = QPushButton("Manage Profiles…")
        btn_manage_profiles.clicked.connect(self._open_project_profiles_dialog)
        profile_bar.addWidget(btn_manage_profiles)
        root.addLayout(profile_bar)

        self._prompt_model_tabs = QTabWidget()
        self._prompt_model_layouts = {}
        self._prompt_profile_editors = {}
        for model_key, title in (("schnell", "Schnell"), ("dev", "Dev"), ("flux2", "FLUX.2")):
            model_page = QWidget()
            model_layout = QVBoxLayout(model_page)
            profile_label = QLabel("Qwen instructions")
            profile_label.setStyleSheet("font-weight:bold;")
            profile_editor = QPlainTextEdit()
            profile_editor.setMaximumHeight(110)
            profile_editor.setPlaceholderText("Instructions used by Llama for this model")
            self._prompt_profile_editors[model_key] = profile_editor
            btn_save_profile = QPushButton("Save Instructions")
            btn_save_profile.clicked.connect(
                lambda _checked, key=model_key: self._save_prompt_profile(key)
            )
            btn_regenerate = QPushButton("Regenerate This Model")
            btn_regenerate.clicked.connect(
                lambda _checked, key=model_key: self._regenerate_prompt_model(key)
            )
            profile_row = QHBoxLayout()
            profile_row.addWidget(profile_label)
            profile_row.addStretch(1)
            profile_row.addWidget(btn_save_profile)
            profile_row.addWidget(btn_regenerate)
            model_layout.addLayout(profile_row)
            model_layout.addWidget(profile_editor)

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            cards_widget = QWidget()
            cards_layout = QVBoxLayout(cards_widget)
            cards_layout.setSpacing(8)
            cards_layout.addStretch(1)
            self._prompt_model_layouts[model_key] = cards_layout
            scroll.setWidget(cards_widget)
            model_layout.addWidget(scroll, 1)
            self._prompt_model_tabs.addTab(model_page, title)
        root.addWidget(self._prompt_model_tabs, 1)

        actions = QHBoxLayout()
        actions.addStretch(1)
        btn_build = QPushButton("Build Prompts")
        btn_build.setToolTip("Generate missing Llama prompts from each scene's English text")
        btn_build.clicked.connect(self._prompts_build_all)
        btn_reload = QPushButton("Reload from Dubbing")
        btn_reload.setToolTip("Reload Italian translations from output/dubbing.yaml")
        btn_reload.clicked.connect(self._prompts_reload_from_dubbing)
        actions.addWidget(btn_build)
        actions.addWidget(btn_reload)
        root.addLayout(actions)
        return page

    def _prompt_profiles_dir(self) -> str:
        return os.path.join(os.path.dirname(os.path.dirname(__file__)), "config", "prompt_profiles")

    def _project_profiles_config_dir(self) -> str:
        return os.path.join(os.path.dirname(os.path.dirname(__file__)), "config")

    def _project_profile_settings_path(self) -> str:
        project = self.project_path_input.text().strip()
        return os.path.join(project, "output", "project_profile.yaml")

    def _load_project_profile_key(self) -> str:
        import yaml as _yaml
        from prompts.project_profiles import NONE_KEY

        path = self._project_profile_settings_path()
        if not os.path.exists(path):
            return NONE_KEY
        try:
            data = _yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
            return str(data.get("selected_profile_key", NONE_KEY)) or NONE_KEY
        except Exception:
            return NONE_KEY

    def _save_project_profile_key(self, key: str) -> None:
        import yaml as _yaml

        project = self.project_path_input.text().strip()
        if not project:
            return
        path = self._project_profile_settings_path()
        os.makedirs(os.path.dirname(path), exist_ok=True)
        Path(path).write_text(
            _yaml.safe_dump({"selected_profile_key": key}, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )

    def _populate_project_profile_combo(self) -> None:
        from prompts.project_profiles import NONE_KEY, NONE_LABEL, load_project_profiles

        if not hasattr(self, "_project_profile_combo"):
            return
        profiles = load_project_profiles(self._project_profiles_config_dir())
        combo = self._project_profile_combo
        combo.blockSignals(True)
        combo.clear()
        combo.addItem(NONE_LABEL, NONE_KEY)
        for key in sorted(profiles.keys()):
            combo.addItem(key.replace("_", " ").title(), key)
        selected = self._load_project_profile_key()
        index = combo.findData(selected)
        combo.setCurrentIndex(index if index >= 0 else 0)
        combo.blockSignals(False)

    def _on_project_profile_selected(self, _index: int) -> None:
        key = self._project_profile_combo.currentData()
        self._save_project_profile_key(str(key or "none"))

    def _open_project_profiles_dialog(self) -> None:
        from PyQt6.QtWidgets import QDialog, QListWidget, QListWidgetItem, QCheckBox
        from prompts.project_profiles import (
            load_project_profiles, save_project_profiles, profiles_path,
        )

        config_dir = self._project_profiles_config_dir()
        profiles = load_project_profiles(config_dir)

        dialog = QDialog(self)
        dialog.setWindowTitle("Manage Project Aesthetic Profiles")
        dialog.resize(900, 600)
        dialog.setStyleSheet(
            "QListWidget, QLineEdit, QPlainTextEdit, QLabel, QPushButton, QCheckBox { font-size:14px; }"
        )
        layout = QHBoxLayout(dialog)

        list_widget = QListWidget()
        for key in sorted(profiles.keys()):
            list_widget.addItem(QListWidgetItem(key))
        layout.addWidget(list_widget, 1)

        right = QVBoxLayout()

        chk_show_yaml = QCheckBox("Show YAML source")
        right.addWidget(chk_show_yaml)

        editor_label = QLabel("Profile key (letters, digits, underscore):")
        right.addWidget(editor_label)
        key_input = QLineEdit()
        right.addWidget(key_input)
        text_label = QLabel("Steering text (appended after each model's Qwen instructions):")
        right.addWidget(text_label)
        text_editor = QPlainTextEdit()
        right.addWidget(text_editor, 1)

        yaml_viewer = QPlainTextEdit()
        yaml_viewer.setReadOnly(True)
        yaml_viewer.setFont(QFont("Consolas", 11))
        yaml_viewer.hide()
        right.addWidget(yaml_viewer, 1)

        yaml_buttons_row = QHBoxLayout()
        btn_edit_yaml = QPushButton("Edit YAML")
        btn_save_yaml = QPushButton("Save YAML")
        btn_edit_yaml.hide()
        btn_save_yaml.hide()
        btn_save_yaml.setEnabled(False)
        yaml_buttons_row.addWidget(btn_edit_yaml)
        yaml_buttons_row.addWidget(btn_save_yaml)
        right.addLayout(yaml_buttons_row)

        def _refresh_yaml_viewer() -> None:
            try:
                yaml_viewer.setPlainText(Path(profiles_path(config_dir)).read_text(encoding="utf-8"))
            except Exception as exc:
                yaml_viewer.setPlainText(f"# Unable to read YAML file: {exc}")

        def _toggle_yaml_view(checked: bool) -> None:
            key_input.setEnabled(not checked)
            text_editor.setEnabled(not checked)
            editor_label.setVisible(not checked)
            text_label.setVisible(not checked)
            text_editor.setVisible(not checked)
            yaml_viewer.setVisible(checked)
            btn_edit_yaml.setVisible(checked)
            btn_save_yaml.setVisible(checked)
            yaml_viewer.setReadOnly(True)
            btn_save_yaml.setEnabled(False)
            if checked:
                _refresh_yaml_viewer()

        def _enable_yaml_edit() -> None:
            yaml_viewer.setReadOnly(False)
            btn_save_yaml.setEnabled(True)
            yaml_viewer.setFocus()

        def _save_yaml() -> None:
            import yaml as _yaml

            raw = yaml_viewer.toPlainText()
            try:
                parsed = _yaml.safe_load(raw) or {}
            except Exception as exc:
                QMessageBox.warning(dialog, "Invalid YAML", f"Could not parse YAML:\n{exc}")
                return
            if not isinstance(parsed, dict) or not all(isinstance(v, str) for v in parsed.values()):
                QMessageBox.warning(dialog, "Invalid YAML", "YAML must be a mapping of profile_key: text.")
                return
            profiles.clear()
            profiles.update(parsed)
            save_project_profiles(config_dir, profiles)
            list_widget.clear()
            for existing_key in sorted(profiles.keys()):
                list_widget.addItem(QListWidgetItem(existing_key))
            key_input.clear()
            text_editor.clear()
            self._populate_project_profile_combo()
            yaml_viewer.setReadOnly(True)
            btn_save_yaml.setEnabled(False)
            _refresh_yaml_viewer()

        chk_show_yaml.toggled.connect(_toggle_yaml_view)
        btn_edit_yaml.clicked.connect(_enable_yaml_edit)
        btn_save_yaml.clicked.connect(_save_yaml)

        buttons_row = QHBoxLayout()
        btn_new = QPushButton("New")
        btn_save = QPushButton("Save")
        btn_delete = QPushButton("Delete")
        buttons_row.addWidget(btn_new)
        buttons_row.addWidget(btn_save)
        buttons_row.addWidget(btn_delete)
        right.addLayout(buttons_row)

        close_row = QHBoxLayout()
        close_row.addStretch(1)
        btn_close = QPushButton("Close")
        close_row.addWidget(btn_close)
        right.addLayout(close_row)
        layout.addLayout(right, 1)

        def _load_selected(item) -> None:
            if item is None:
                key_input.clear()
                text_editor.clear()
                return
            key = item.text()
            key_input.setText(key)
            text_editor.setPlainText(profiles.get(key, ""))

        def _new() -> None:
            list_widget.clearSelection()
            key_input.clear()
            text_editor.clear()
            key_input.setFocus()

        def _save() -> None:
            key = key_input.text().strip().lower().replace(" ", "_")
            if not key or not re.match(r"^[a-z0-9_]+$", key):
                QMessageBox.warning(dialog, "Invalid key", "Use lowercase letters, digits, and underscores only.")
                return
            profiles[key] = text_editor.toPlainText().strip()
            save_project_profiles(config_dir, profiles)
            list_widget.clear()
            for existing_key in sorted(profiles.keys()):
                list_widget.addItem(QListWidgetItem(existing_key))
            self._populate_project_profile_combo()
            if chk_show_yaml.isChecked():
                _refresh_yaml_viewer()

        def _delete() -> None:
            item = list_widget.currentItem()
            if item is None:
                return
            key = item.text()
            answer = QMessageBox.question(dialog, "Delete profile", f"Delete profile '{key}'?")
            if answer != QMessageBox.StandardButton.Yes:
                return
            profiles.pop(key, None)
            save_project_profiles(config_dir, profiles)
            list_widget.takeItem(list_widget.row(item))
            key_input.clear()
            text_editor.clear()
            self._populate_project_profile_combo()
            if chk_show_yaml.isChecked():
                _refresh_yaml_viewer()

        list_widget.currentItemChanged.connect(lambda current, _prev: _load_selected(current))
        btn_new.clicked.connect(_new)
        btn_save.clicked.connect(_save)
        btn_delete.clicked.connect(_delete)
        btn_close.clicked.connect(dialog.accept)
        dialog.exec()

    def _prompts_build_all(self) -> None:
        path = self.project_path_input.text().strip()
        if path:
            self.controller.set_project_path(path)
        key = self._project_profile_combo.currentData() if hasattr(self, "_project_profile_combo") else None
        self.controller.run_pipeline("prompts", {"project_profile_key": str(key or "none")})

    def _save_prompt_profile(self, model_key: str) -> None:
        import yaml as _yaml

        path = os.path.join(self._prompt_profiles_dir(), f"{model_key}.yaml")
        try:
            profile = _yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
            profile["system_instruction"] = self._prompt_profile_editors[model_key].toPlainText().strip()
            Path(path).write_text(
                _yaml.safe_dump(profile, allow_unicode=True, sort_keys=False), encoding="utf-8"
            )
            self._prompts_status_label.setText(f"Saved {model_key} Qwen instructions.")
        except Exception as exc:
            QMessageBox.critical(self, "Save instructions failed", str(exc))

    def _regenerate_prompt_model(self, model_key: str) -> None:
        self._save_prompt_profile(model_key)
        answer = QMessageBox.question(
            self,
            "Regenerate model prompts",
            f"Replace all {model_key} prompts, including manual edits, using the current instructions?",
        )
        if answer != QMessageBox.StandardButton.Yes:
            return
        path = self.project_path_input.text().strip()
        if path:
            self.controller.set_project_path(path)
        key = self._project_profile_combo.currentData() if hasattr(self, "_project_profile_combo") else None
        self.controller.run_pipeline("prompts", {
            "prompt_model_key": model_key,
            "force_regenerate_prompts": True,
            "project_profile_key": str(key or "none"),
        })

    def _prompts_reload_from_dubbing(self) -> None:
        if getattr(self, "_dub_editors", None):
            self._dub_save()
        self._prompts_refresh()

    def _prompts_project_data(self):
        import yaml as _yaml

        project = self.project_path_input.text().strip()
        if not project:
            return [], {}, {}, {}

        output = os.path.join(project, "output")

        def _load(name: str):
            path = os.path.join(output, name)
            if not os.path.exists(path):
                return {}
            try:
                return _yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
            except Exception:
                return {}

        scenes = _load("scenes.yaml").get("scenes", [])
        return scenes, _load("dubbing.yaml"), _load("prompts.yaml"), _load("prompt_overrides.yaml")

    def _prompts_refresh(self) -> None:
        if not hasattr(self, "_prompt_model_layouts"):
            return
        self._populate_project_profile_combo()
        for cards_layout in self._prompt_model_layouts.values():
            while cards_layout.count() > 1:
                item = cards_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()

        scenes, dubbing, prompts, overrides = self._prompts_project_data()
        if not scenes:
            self._prompts_status_label.setText("No scenes found. Split scenes from the Script tab first.")
            return

        import yaml as _yaml
        project = self.project_path_input.text().strip()
        model_prompts_path = os.path.join(project, "output", "model_prompts.yaml")
        try:
            model_prompts = _yaml.safe_load(Path(model_prompts_path).read_text(encoding="utf-8")) or {}
        except Exception:
            model_prompts = {}

        saved_preview_beats = set()
        draft_dir = os.path.join(project, "output", "draft")
        if os.path.isdir(draft_dir):
            for filename in os.listdir(draft_dir):
                match = re.match(
                    r"^scene_(\d+)_schnell_b(\d+)_v\d+\.(?:png|jpg|jpeg)$",
                    filename,
                    re.IGNORECASE,
                )
                if match:
                    saved_preview_beats.add((int(match.group(1)), int(match.group(2))))
                    continue
                legacy_match = re.match(
                    r"^scene_(\d+)\.(?:png|jpg|jpeg)$", filename, re.IGNORECASE)
                if legacy_match:
                    saved_preview_beats.add((int(legacy_match.group(1)), 1))

        saved_lightbox_beats = set()
        lightbox_dir = os.path.join(project, "output", "lightbox")
        if os.path.isdir(lightbox_dir):
            for filename in os.listdir(lightbox_dir):
                match = re.match(
                    r"^scene_(\d+)_(dev|flux2)_b(\d+)_v\d+\.(?:png|jpg|jpeg)$",
                    filename,
                    re.IGNORECASE,
                )
                if match:
                    saved_lightbox_beats.add(
                        (match.group(2).lower(), int(match.group(1)), int(match.group(3)))
                    )
                    continue
                legacy_match = re.match(
                    r"^scene_(\d+)_(dev|flux2)_v\d+\.(?:png|jpg|jpeg)$",
                    filename,
                    re.IGNORECASE,
                )
                if legacy_match:
                    saved_lightbox_beats.add(
                        (legacy_match.group(2).lower(), int(legacy_match.group(1)), 1)
                    )

        for model_key, editor in self._prompt_profile_editors.items():
            try:
                profile_path = os.path.join(self._prompt_profiles_dir(), f"{model_key}.yaml")
                profile = _yaml.safe_load(Path(profile_path).read_text(encoding="utf-8")) or {}
                editor.setPlainText(str(profile.get("system_instruction", "")))
            except Exception:
                editor.clear()

        for model_key, cards_layout in self._prompt_model_layouts.items():
            for scene in scenes:
                sid = int(scene["id"])
                scene_entry = model_prompts.get(sid) or model_prompts.get(str(sid)) or {}
                models = scene_entry.get("models", {}) if isinstance(scene_entry, dict) else {}
                model_entry = models.get(model_key, {})
                rows = model_entry.get("prompts", []) if isinstance(model_entry, dict) else []

                card = QWidget()
                card.setObjectName("promptCard")
                card.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
                card.setToolTip("Click to edit this model's visual beats")
                card.setStyleSheet(
                    "QWidget#promptCard { background:#1D1B20; border:1px solid #36343B; border-radius:4px; }"
                    "QWidget#promptCard:hover { border:1px solid #96BDE2; }"
                )
                layout = QVBoxLayout(card)
                layout.setContentsMargins(10, 8, 10, 8)
                layout.setSpacing(4)

                manual = any(row.get("source") == "manually_edited" for row in rows if isinstance(row, dict))
                title = QLabel(f"Scene {sid} | {len(rows)} visual beat(s)" + (" | manual" if manual else ""))
                title.setStyleSheet("color:#96BDE2; font-weight:bold; font-size:11px; border:none;")
                title.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents)
                layout.addWidget(title)
                for index, row in enumerate(rows, 1):
                    text = str(row.get("text", "")) if isinstance(row, dict) else str(row)
                    beat_index = int(row.get("beat", index)) if isinstance(row, dict) else index
                    has_preview = model_key == "schnell" and (sid, beat_index) in saved_preview_beats
                    has_lightbox = (
                        model_key in ("dev", "flux2")
                        and (model_key, sid, beat_index) in saved_lightbox_beats
                    )
                    has_saved_image = has_preview or has_lightbox
                    value_label = QLabel(f"{index}. {text}")
                    value_label.setWordWrap(True)
                    value_label.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
                    if has_saved_image:
                        saved_location = "Preview" if has_preview else "Lightbox"
                        value_label.setToolTip(
                            f"{saved_location} image saved. Click to open this visual beat")
                        value_label.setStyleSheet(
                            "color:#C9A6E6; font-size:12px; background:#173326; "
                            "border:2px solid #9FD6B8; border-radius:2px; padding:5px;"
                        )
                    else:
                        value_label.setToolTip("Open this visual beat")
                        value_label.setStyleSheet(
                            "color:#C9A6E6; font-size:12px; background:#131118; "
                            "border:1px solid #2a2830; border-radius:2px; padding:5px;"
                        )
                    value_label.mousePressEvent = (
                        lambda event, s=sid, key=model_key, beat=beat_index:
                        self._open_prompt_beat_dialog(s, key, beat)
                    )
                    layout.addWidget(value_label)
                if not rows:
                    layout.addWidget(QLabel("Not generated"))

                cards_layout.insertWidget(cards_layout.count() - 1, card)

        self._prompts_status_label.setText(
            f"{len(scenes)} scene(s). Build Prompts generates all three profiles; click a card to edit beats."
        )

    def _edit_model_prompts(self, scene_id: int, model_key: str) -> None:
        import yaml as _yaml
        from PyQt6.QtWidgets import QDialog

        project = self.project_path_input.text().strip()
        path = os.path.join(project, "output", "model_prompts.yaml")
        try:
            data = _yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
        except Exception:
            data = {}
        scene_entry = data.get(scene_id) or data.get(str(scene_id)) or {"scene_id": scene_id, "models": {}}
        models = scene_entry.setdefault("models", {})
        model_entry = models.setdefault(model_key, {"profile": f"{model_key}.yaml", "prompts": []})
        rows = model_entry.get("prompts", [])

        dialog = QDialog(self)
        dialog.setWindowTitle(f"Scene {scene_id} | {model_key} prompts")
        dialog.resize(820, 560)
        layout = QVBoxLayout(dialog)
        layout.addWidget(QLabel("Separate visual beats with a blank line."))
        editor = QPlainTextEdit()
        editor.setPlainText("\n\n".join(str(row.get("text", "")) for row in rows if isinstance(row, dict)))
        layout.addWidget(editor, 1)
        actions = QHBoxLayout()
        actions.addStretch(1)
        cancel = QPushButton("Close")
        save = QPushButton("Save Manual Prompts")
        actions.addWidget(cancel)
        actions.addWidget(save)
        layout.addLayout(actions)

        def _save():
            texts = [part.strip() for part in editor.toPlainText().split("\n\n") if part.strip()]
            model_entry["prompts"] = [
                {
                    "id": f"scene_{scene_id:03d}_beat_{index:02d}_{model_key}",
                    "beat": index,
                    "text": text,
                    "source": "manually_edited",
                }
                for index, text in enumerate(texts, 1)
            ]
            data[scene_id] = scene_entry
            data.pop(str(scene_id), None)
            os.makedirs(os.path.dirname(path), exist_ok=True)
            Path(path).write_text(
                _yaml.safe_dump(data, allow_unicode=True, sort_keys=False), encoding="utf-8"
            )
            if model_key == "schnell" and texts:
                legacy_path = os.path.join(project, "output", "prompts.yaml")
                try:
                    legacy = _yaml.safe_load(Path(legacy_path).read_text(encoding="utf-8")) or {}
                except Exception:
                    legacy = {}
                legacy[scene_id] = texts[0]
                legacy.pop(str(scene_id), None)
                Path(legacy_path).write_text(
                    _yaml.safe_dump(legacy, allow_unicode=True, sort_keys=False), encoding="utf-8"
                )
            dialog.accept()
            self._prompts_refresh()

        cancel.clicked.connect(dialog.reject)
        save.clicked.connect(_save)
        dialog.exec()

    def _open_prompt_beat_dialog(self, scene_id: int, model_key: str, beat_index: int) -> None:
        import yaml as _yaml
        from PyQt6.QtWidgets import QDialog, QTextEdit
        from prompts.model_prompt_service import ModelPromptService

        project = self.project_path_input.text().strip()
        prompts_path = os.path.join(project, "output", "model_prompts.yaml")
        scenes, _dubbing, _prompts, _overrides = self._prompts_project_data()
        scene = next((item for item in scenes if int(item["id"]) == scene_id), {})
        try:
            data = _yaml.safe_load(Path(prompts_path).read_text(encoding="utf-8")) or {}
            scene_entry = data.get(scene_id) or data.get(str(scene_id)) or {}
            model_entry = scene_entry.get("models", {}).get(model_key, {})
            rows = model_entry.get("prompts", [])
            # Match by the row's own "beat" field first — list position can drift
            # from the beat number once rows are regenerated/edited out of order.
            row = next(
                (r for r in rows if isinstance(r, dict) and int(r.get("beat", 0)) == beat_index),
                None,
            )
            if row is None:
                row = rows[beat_index - 1]
        except Exception as exc:
            QMessageBox.warning(self, "Prompt unavailable", str(exc))
            return

        visual_beat = str(row.get("visual_beat") or scene.get("text", ""))
        generated_prompt = str(row.get("generated_prompt") or row.get("text", ""))
        lightbox_dir = os.path.join(project, "output", "lightbox")
        preview_dir = (
            os.path.join(project, "output", "draft")
            if model_key == "schnell" else lightbox_dir
        )
        preview_candidates = [
            os.path.join(preview_dir, f"scene_{scene_id:03d}_{model_key}_b{beat_index:02d}_v{variant}.png")
            for variant in (2, 1, 3)
        ]
        if model_key == "schnell" and beat_index == 1:
            preview_candidates.append(os.path.join(preview_dir, f"scene_{scene_id:03d}.png"))
        image_path = next((path for path in preview_candidates if os.path.exists(path)), "")

        dialog = QDialog(self)
        dialog.setWindowTitle(f"Scene {scene_id} | {model_key} | Beat {beat_index}")
        dialog.resize(900, 820)
        layout = QVBoxLayout(dialog)

        popup_controls = QHBoxLayout()
        popup_controls.addStretch(1)
        popup_controls.addWidget(QLabel("Text size:"))
        font_slider = QSlider(Qt.Orientation.Horizontal)
        font_slider.setRange(8, 36)
        font_slider.setValue(15)
        font_slider.setFixedWidth(120)
        font_slider.setToolTip("Adjust popup text size")
        popup_controls.addWidget(font_slider)
        font_size_label = QLabel("15px")
        font_size_label.setFixedWidth(32)
        popup_controls.addWidget(font_size_label)
        fullscreen_button = QPushButton("⛶")
        fullscreen_button.setFixedSize(30, 30)
        fullscreen_button.setToolTip("Toggle full screen")
        popup_controls.addWidget(fullscreen_button)
        layout.addLayout(popup_controls)

        preview = QLabel("No image has been generated for this beat")
        preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        preview.setMinimumHeight(220)
        if image_path:
            pixmap = QPixmap(image_path)
            preview.setPixmap(pixmap.scaled(
                840, 300, Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation,
            ))
            preview.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
            preview.mousePressEvent = lambda event: self._open_image_viewer(image_path)
        layout.addWidget(preview)

        popup_text_editors = []

        def _readonly_section(title: str, text: str, height: int, color: str):
            title_label = QLabel(title)
            title_label.setStyleSheet(f"color:{color}; font-weight:bold;")
            layout.addWidget(title_label)
            editor = QTextEdit()
            editor.setPlainText(text)
            editor.setReadOnly(True)
            editor.setFixedHeight(height)
            palette = editor.palette()
            palette.setColor(editor.backgroundRole(), QColor("#131118"))
            palette.setColor(editor.foregroundRole(), QColor(color))
            palette.setColor(palette.ColorRole.Base, QColor("#131118"))
            palette.setColor(palette.ColorRole.Text, QColor(color))
            editor.setPalette(palette)
            layout.addWidget(editor)
            popup_text_editors.append((editor, color, height))

        _readonly_section("Original script", str(scene.get("text", "")), 75, "#96BDE2")
        _readonly_section("Visual beat identified by Qwen", visual_beat, 75, "#9FD6B8")
        _readonly_section("Original prompt generated by Qwen", generated_prompt, 105, "#C9A6E6")

        layout.addWidget(QLabel("Working prompt"))
        prompt_editor = QTextEdit()
        prompt_editor.setPlainText(str(row.get("text", "")))
        prompt_editor.setMinimumHeight(130)
        layout.addWidget(prompt_editor)
        popup_text_editors.append((prompt_editor, "#E6E1E5", None))

        def _set_popup_font_size(size: int):
            font_size_label.setText(f"{size}px")
            for editor, color, _normal_height in popup_text_editors:
                font = editor.font()
                font.setPixelSize(size)
                text_format = QTextCharFormat()
                text_format.setFont(font)
                text_format.setForeground(QColor(color))
                cursor = editor.textCursor()
                cursor_position = cursor.position()
                cursor.select(QTextCursor.SelectionType.Document)
                cursor.mergeCharFormat(text_format)
                cursor.clearSelection()
                cursor.setPosition(min(cursor_position, len(editor.toPlainText())))
                editor.setTextCursor(cursor)
                editor.setCurrentCharFormat(text_format)

        def _set_prompt_sections_expanded(expanded: bool):
            for editor, _color, normal_height in popup_text_editors:
                if expanded:
                    editor.setMinimumHeight(75)
                    editor.setMaximumHeight(16777215)
                    layout.setStretchFactor(editor, 1)
                elif normal_height is not None:
                    editor.setFixedHeight(normal_height)
                    layout.setStretchFactor(editor, 0)
                else:
                    editor.setMinimumHeight(130)
                    editor.setMaximumHeight(16777215)
                    layout.setStretchFactor(editor, 1)

        def _toggle_fullscreen():
            if dialog.isFullScreen():
                _set_prompt_sections_expanded(False)
                dialog.showNormal()
                fullscreen_button.setText("⛶")
                fullscreen_button.setToolTip("Enter full screen")
            else:
                _set_prompt_sections_expanded(True)
                dialog.showFullScreen()
                fullscreen_button.setText("▣")
                fullscreen_button.setToolTip("Exit full screen")

        font_slider.valueChanged.connect(_set_popup_font_size)
        fullscreen_button.clicked.connect(_toggle_fullscreen)
        _set_popup_font_size(font_slider.value())

        status = QLabel("")
        layout.addWidget(status)
        progress = QProgressBar()
        progress.setRange(0, 100)
        progress.setVisible(False)
        layout.addWidget(progress)
        actions = QHBoxLayout()
        actions.addStretch(1)
        cancel = QPushButton("Close")
        regenerate = QPushButton("Regenerate with Qwen")
        save = QPushButton("Save Prompt")
        generate_image = QPushButton("Generate Image")
        save_image = QPushButton(
            "Save in Preview" if model_key == "schnell" else "Save in Lightbox")
        save_image.setEnabled(False)
        actions.addWidget(cancel)
        actions.addWidget(regenerate)
        actions.addWidget(save)
        actions.addWidget(generate_image)
        actions.addWidget(save_image)
        layout.addLayout(actions)
        candidate_state = {"path": ""}
        saved_prompt_state = {"text": prompt_editor.toPlainText().strip()}

        def _set_save_button_style(button: QPushButton, needs_save: bool):
            if needs_save:
                button.setStyleSheet(
                    "QPushButton { background:#5B2026; color:#FFD7D9; border:1px solid #E73A4B; }"
                    "QPushButton:hover { background:#7A2B33; }"
                )
            else:
                button.setStyleSheet("")

        def _update_save_states():
            prompt_changed = prompt_editor.toPlainText().strip() != saved_prompt_state["text"]
            _set_save_button_style(save, prompt_changed)
            _set_save_button_style(save_image, bool(candidate_state["path"]))

        prompt_editor.textChanged.connect(_update_save_states)
        _update_save_states()

        def _regenerate():
            from prompts.project_profiles import get_profile_text, load_project_profiles

            project_key = self._load_project_profile_key()
            project_profile_text = get_profile_text(
                load_project_profiles(self._project_profiles_config_dir()), project_key
            )
            service = ModelPromptService(
                self._prompt_profiles_dir(),
                self.ollama_model_input.text().strip() or "qwen3:8b",
                self.ollama_host_input.text().strip() or "http://localhost:11434",
                project_profile_text=project_profile_text,
            )
            regenerate.setEnabled(False)
            status.setText("Regenerating this prompt...")
            try:
                replacement = service.regenerate_prompt(scene, model_key, visual_beat)
                prompt_editor.setPlainText(replacement)
                # Do NOT touch row["generated_prompt"] here — it must keep recording
                # Qwen's true original output, not the latest regeneration.
                status.setText("New Qwen prompt ready. Save to keep it.")
            except Exception as exc:
                QMessageBox.warning(dialog, "Llama regeneration failed", str(exc))
                status.clear()
            finally:
                regenerate.setEnabled(True)

        def _save(close_dialog: bool = True):
            text = prompt_editor.toPlainText().strip()
            if not text:
                QMessageBox.warning(dialog, "Empty prompt", "Enter a prompt before saving.")
                return False
            previous_text = str(row.get("text", "")).strip()
            row["text"] = text
            row["source"] = "manually_edited"
            data[scene_id] = scene_entry
            data.pop(str(scene_id), None)
            Path(prompts_path).write_text(
                _yaml.safe_dump(data, allow_unicode=True, sort_keys=False), encoding="utf-8"
            )
            saved_prompt_state["text"] = text
            _update_save_states()
            if model_key == "schnell" and beat_index == 1:
                legacy_path = os.path.join(project, "output", "prompts.yaml")
                try:
                    legacy = _yaml.safe_load(Path(legacy_path).read_text(encoding="utf-8")) or {}
                except Exception:
                    legacy = {}
                legacy[scene_id] = text
                legacy.pop(str(scene_id), None)
                Path(legacy_path).write_text(
                    _yaml.safe_dump(legacy, allow_unicode=True, sort_keys=False), encoding="utf-8"
                )
            if text != previous_text:
                invalidated = {
                    f"scene_{scene_id:03d}_{model_key}_b{beat_index:02d}_v{variant}.png"
                    for variant in (1, 2, 3)
                }
                for filename in invalidated:
                    path = os.path.join(lightbox_dir, filename)
                    if os.path.exists(path):
                        os.remove(path)
                selections_path = os.path.join(project, "output", "lightbox_selections.yaml")
                if os.path.exists(selections_path):
                    selections = _yaml.safe_load(
                        Path(selections_path).read_text(encoding="utf-8")) or {}
                    key = scene_id if scene_id in selections else str(scene_id)
                    if key in selections and isinstance(selections[key], list):
                        selections[key] = [name for name in selections[key] if name not in invalidated]
                        if not selections[key]:
                            selections.pop(key)
                        Path(selections_path).write_text(
                            _yaml.safe_dump(selections, allow_unicode=True, sort_keys=False),
                            encoding="utf-8",
                        )
            if close_dialog:
                dialog.accept()
                self._prompts_refresh()
                self._refresh_lightbox()
            else:
                status.setText("Prompt saved.")
            return True

        def _set_image_running(running: bool):
            generate_image.setEnabled(not running)
            regenerate.setEnabled(not running)
            save.setEnabled(not running)
            cancel.setEnabled(not running)
            save_image.setEnabled(not running and bool(candidate_state["path"]))
            progress.setVisible(running)

        # Defined once (not re-created per click) so a stray connection left over from
        # a previous click, or from this dialog being closed mid-generation, can always
        # be found and disconnected by identity instead of silently piling up on the
        # controller's shared signals — each extra stray connection would otherwise fire
        # against this dialog's widgets even after they were deleted, raising a
        # RuntimeError deep inside Qt's signal dispatch and stalling the event loop.
        def _show_progress(value: int, message: str):
            if sip.isdeleted(dialog):
                return
            progress.setValue(value)
            status.setText(message)

        def _show_candidate(success: bool, payload: str):
            try:
                self.controller.pipeline_finished.disconnect(_show_candidate)
                self.controller.pipeline_progress.disconnect(_show_progress)
            except Exception:
                pass
            if sip.isdeleted(dialog):
                return
            _set_image_running(False)
            if not success:
                status.setText(f"Image generation failed: {payload}")
                return
            candidate_state["path"] = payload
            save_image.setEnabled(True)
            _update_save_states()
            pixmap = QPixmap(payload)
            if not pixmap.isNull():
                preview.setPixmap(pixmap.scaled(
                    840, 300, Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation,
                ))
                preview.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
                preview.mousePressEvent = lambda event: self._open_image_viewer(payload)
            progress.setValue(100)
            status.setText("Candidate ready. Save it or adjust the prompt and generate again.")

        def _disconnect_generation_signals():
            try:
                self.controller.pipeline_finished.disconnect(_show_candidate)
            except Exception:
                pass
            try:
                self.controller.pipeline_progress.disconnect(_show_progress)
            except Exception:
                pass

        dialog.finished.connect(lambda _result: _disconnect_generation_signals())

        def _generate_image():
            prompt = prompt_editor.toPlainText().strip()
            if not prompt:
                QMessageBox.warning(dialog, "Empty prompt", "Enter a prompt before generating an image.")
                return
            if getattr(self.controller, "_thread", None) is not None:
                QMessageBox.warning(dialog, "Pipeline busy", "Wait for the current pipeline operation to finish.")
                return
            # Guard against a leftover connection from a previous click in this same
            # dialog (e.g. Generate Image clicked twice before the first finished).
            _disconnect_generation_signals()
            _set_image_running(True)
            progress.setValue(0)
            status.setText(f"Generating neutral-seed {model_key} image...")
            self.controller.pipeline_progress.connect(_show_progress)
            self.controller.pipeline_finished.connect(_show_candidate)
            self.controller.run_pipeline("prompt_image_candidate", {
                "preview_scene_id": scene_id,
                "preview_beat_index": beat_index,
                "preview_model_key": model_key,
                "preview_prompt": prompt,
            })

        def _save_image():
            candidate_path = candidate_state["path"]
            if not candidate_path or not os.path.exists(candidate_path):
                QMessageBox.warning(dialog, "No candidate", "Generate an image first.")
                return
            if not _save(close_dialog=False):
                return
            if model_key == "schnell":
                destination_dir = os.path.join(project, "output", "draft")
            else:
                destination_dir = lightbox_dir
            os.makedirs(destination_dir, exist_ok=True)
            destination = os.path.join(
                destination_dir,
                f"scene_{scene_id:03d}_{model_key}_b{beat_index:02d}_v2.png",
            )
            import shutil as _shutil
            _shutil.copy2(candidate_path, destination)
            os.remove(candidate_path)
            candidate_state["path"] = ""
            save_image.setEnabled(False)
            _update_save_states()
            preview.mousePressEvent = lambda event: self._open_image_viewer(destination)
            status.setText(f"Saved neutral image to {destination_dir}.")
            self._prompts_refresh()
            self._refresh_draft_grid()
            self._refresh_lightbox()

        cancel.clicked.connect(dialog.reject)
        regenerate.clicked.connect(_regenerate)
        save.clicked.connect(lambda: _save(close_dialog=False))
        generate_image.clicked.connect(_generate_image)
        save_image.clicked.connect(_save_image)
        dialog.exec()

    def _open_prompt_scene_dialog(self, scene_id: int, candidate: bool = False,
                                  prior_override=None) -> None:
        import yaml as _yaml
        from PyQt6.QtWidgets import QDialog, QTextEdit

        scenes, dubbing, prompts, overrides = self._prompts_project_data()
        scene = next((s for s in scenes if int(s["id"]) == scene_id), {})
        original = scene.get("text", "")
        dub_entry = dubbing.get(scene_id) or dubbing.get(str(scene_id)) or {}
        italian = dub_entry.get("dubbed", "") if isinstance(dub_entry, dict) else str(dub_entry)
        llama_prompt = prompts.get(scene_id) or prompts.get(str(scene_id)) or ""
        working_prompt = overrides.get(scene_id) or overrides.get(str(scene_id)) or llama_prompt

        project = self.project_path_input.text().strip()
        image_path = os.path.join(project, "output", "draft", f"scene_{scene_id:03d}.png")
        backup_path = image_path + ".prompt_backup"
        overrides_path = os.path.join(project, "output", "prompt_overrides.yaml")

        dlg = QDialog(self)
        dlg.setWindowTitle(f"Scene {scene_id} prompt refinement")
        dlg.setModal(True)
        dlg.resize(900, 760)
        dlg.setMinimumSize(720, 600)
        dlg.setWindowFlag(Qt.WindowType.WindowMaximizeButtonHint, True)
        layout = QVBoxLayout(dlg)

        preview = QLabel("No preview image yet")
        preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        preview.setMinimumHeight(260)
        if os.path.exists(image_path):
            pixmap = QPixmap(image_path)
            preview.setPixmap(pixmap.scaled(820, 360, Qt.AspectRatioMode.KeepAspectRatio,
                                            Qt.TransformationMode.SmoothTransformation))
            preview.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
            preview.setToolTip("Click to view the full-size image")
            preview.mousePressEvent = lambda event: self._open_image_viewer(image_path)
        layout.addWidget(preview)

        run_status = QLabel("")
        run_status.setStyleSheet("color:#8E8B90; font-size:11px;")
        run_progress = QProgressBar()
        run_progress.setRange(0, 100)
        run_progress.setVisible(False)
        layout.addWidget(run_status)
        layout.addWidget(run_progress)

        for heading, text in (("English original", original), ("Italian translation", italian)):
            heading_label = QLabel(heading)
            heading_label.setStyleSheet("font-size:14px; font-weight:bold;")
            layout.addWidget(heading_label)
            view = QTextEdit()
            view.setPlainText(text)
            view.setReadOnly(True)
            view.setFixedHeight(70)
            view.setStyleSheet("QTextEdit { font-size:16px; }")
            layout.addWidget(view)

        original_prompt_label = QLabel("Original Llama prompt (read only)")
        original_prompt_label.setStyleSheet("font-size:15px; font-weight:bold;")
        layout.addWidget(original_prompt_label)
        original_prompt_view = QTextEdit()
        original_prompt_view.setPlainText(llama_prompt)
        original_prompt_view.setReadOnly(True)
        original_prompt_view.setMinimumHeight(100)
        original_prompt_view.setStyleSheet(
            "QTextEdit { background:#131118; color:#C9A6E6; border:1px solid #36343B; "
            "border-radius:3px; padding:5px; font-size:19px; }"
        )
        layout.addWidget(original_prompt_view)

        working_prompt_label = QLabel("Working prompt")
        working_prompt_label.setStyleSheet("font-size:15px; font-weight:bold;")
        layout.addWidget(working_prompt_label)
        prompt_edit = QTextEdit()
        prompt_edit.setPlainText(working_prompt)
        prompt_edit.setMinimumHeight(120)
        prompt_edit.setStyleSheet("QTextEdit { font-size:19px; }")
        layout.addWidget(prompt_edit)

        tuning = QHBoxLayout()
        tuning.addWidget(QLabel("Qwen model:"))
        llama_model = QLineEdit(self.ollama_model_input.text().strip() or "qwen3:8b")
        tuning.addWidget(llama_model, 1)
        tuning.addWidget(QLabel("Schnell steps:"))
        steps = QSpinBox()
        steps.setRange(1, 20)
        steps.setValue(self.schnell_steps_input.value())
        tuning.addWidget(steps)
        tuning.addWidget(QLabel("Guidance:"))
        guidance = QDoubleSpinBox()
        guidance.setRange(0.0, 10.0)
        guidance.setValue(self.schnell_guidance_input.value())
        tuning.addWidget(guidance)
        layout.addLayout(tuning)

        buttons = QHBoxLayout()
        buttons.addStretch(1)
        btn_close = QPushButton("Discard Candidate" if candidate else "Close")
        btn_llama = QPushButton("Regenerate with Qwen")
        btn_run = QPushButton("Run FLUX Schnell")
        btn_commit = QPushButton("Commit to Preview Images")
        btn_commit.setEnabled(candidate)
        buttons.addWidget(btn_close)
        buttons.addWidget(btn_llama)
        buttons.addWidget(btn_run)
        buttons.addWidget(btn_commit)
        layout.addLayout(buttons)

        def _save_override() -> bool:
            text = prompt_edit.toPlainText().strip()
            if not text:
                QMessageBox.warning(dlg, "Empty prompt", "Enter a prompt first.")
                return False
            overrides[scene_id] = text
            overrides.pop(str(scene_id), None)
            os.makedirs(os.path.dirname(overrides_path), exist_ok=True)
            Path(overrides_path).write_text(
                _yaml.safe_dump(overrides, allow_unicode=True, sort_keys=False), encoding="utf-8"
            )
            self.ollama_model_input.setText(llama_model.text().strip() or "qwen3:8b")
            self.schnell_steps_input.setValue(steps.value())
            self.schnell_guidance_input.setValue(guidance.value())
            self.controller.blockSignals(True)
            try:
                self.controller.save_settings(self._collect_settings_from_form())
            finally:
                self.controller.blockSignals(False)
            return True

        candidate_state = {"resolved": False, "has_candidate": candidate}

        def _set_running(running: bool):
            btn_run.setEnabled(not running)
            btn_llama.setEnabled(not running)
            btn_close.setEnabled(not running)
            btn_commit.setEnabled(not running and candidate_state["has_candidate"])
            run_progress.setVisible(running)

        def _run_preview():
            prompt = prompt_edit.toPlainText().strip()
            if not prompt:
                QMessageBox.warning(dlg, "Empty prompt", "Enter a prompt first.")
                return
            if getattr(self.controller, "_thread", None) is not None:
                QMessageBox.warning(dlg, "Pipeline busy", "Wait for the current pipeline operation to finish.")
                return
            if not candidate_state["has_candidate"] and os.path.exists(image_path):
                if os.path.exists(backup_path):
                    os.remove(backup_path)
                os.replace(image_path, backup_path)

            self.schnell_steps_input.setValue(steps.value())
            self.schnell_guidance_input.setValue(guidance.value())
            self.controller.blockSignals(True)
            try:
                self.controller.save_settings(self._collect_settings_from_form())
            finally:
                self.controller.blockSignals(False)

            _set_running(True)
            run_progress.setValue(0)
            run_status.setText(f"Generating scene {scene_id} with FLUX schnell...")

            def _show_progress(value: int, message: str):
                run_progress.setValue(value)
                run_status.setText(message)

            def _show_candidate(success: bool, payload: str):
                try:
                    self.controller.pipeline_finished.disconnect(_show_candidate)
                    self.controller.pipeline_progress.disconnect(_show_progress)
                except Exception:
                    pass
                _set_running(False)
                if success:
                    candidate_state["has_candidate"] = True
                    btn_close.setText("Discard Candidate")
                    btn_commit.setEnabled(True)
                    pixmap = QPixmap(image_path)
                    preview.setPixmap(pixmap.scaled(
                        820, 360, Qt.AspectRatioMode.KeepAspectRatio,
                        Qt.TransformationMode.SmoothTransformation
                    ))
                    preview.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
                    preview.setToolTip("Click to view the full-size image")
                    preview.mousePressEvent = lambda event: self._open_image_viewer(image_path)
                    run_progress.setValue(100)
                    run_status.setText(f"Scene {scene_id} candidate ready. Commit or adjust and run again.")
                else:
                    if os.path.exists(backup_path):
                        os.replace(backup_path, image_path)
                    run_status.setText(f"Generation failed: {payload}")

            self.controller.pipeline_progress.connect(_show_progress)
            self.controller.pipeline_finished.connect(_show_candidate)
            self.controller.run_pipeline("preview_scene", {
                "preview_scene_id": scene_id,
                "preview_prompt": prompt,
                "schnell_steps": steps.value(),
                "schnell_guidance": guidance.value(),
            })

        def _commit():
            if not _save_override():
                return
            if os.path.exists(backup_path):
                os.remove(backup_path)
            candidate_state["resolved"] = True
            self._prompts_refresh()
            self._refresh_draft_grid()
            dlg.accept()

        def _discard_candidate():
            if candidate_state["has_candidate"] and not candidate_state["resolved"]:
                candidate_state["resolved"] = True
                if os.path.exists(image_path):
                    os.remove(image_path)
                if os.path.exists(backup_path):
                    os.replace(backup_path, image_path)
                self._refresh_draft_grid()
                self._prompts_refresh()

        def _close_or_discard():
            _discard_candidate()
            dlg.reject()

        def _regenerate_with_llama():
            from prompts.prompt_builder import PromptBuilder
            builder = PromptBuilder(
                style_preset=self.style_preset_input.currentText(),
                default_aspect_ratio=self.aspect_ratio_input.currentText(),
                use_ollama=True,
                ollama_model=llama_model.text().strip() or "qwen3:8b",
                ollama_host=self.ollama_host_input.text().strip() or "http://localhost:11434",
            )
            try:
                prompt_edit.setPlainText(builder.build_prompt({"id": scene_id, "text": original}))
            except Exception as exc:
                QMessageBox.warning(dlg, "Llama generation failed", str(exc))

        btn_close.clicked.connect(_close_or_discard)
        dlg.rejected.connect(_discard_candidate)
        btn_llama.clicked.connect(_regenerate_with_llama)
        btn_run.clicked.connect(_run_preview)
        btn_commit.clicked.connect(_commit)

        # Suppress grid rebuilds while this dialog is modal (see _refresh_draft_grid).
        self._draft_dialog_depth = getattr(self, "_draft_dialog_depth", 0) + 1
        try:
            dlg.exec()
        finally:
            self._draft_dialog_depth -= 1
        if self._draft_dialog_depth <= 0 and getattr(self, "_draft_grid_refresh_pending", False):
            self._draft_grid_refresh_pending = False
            self._refresh_draft_grid()

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _dub_yaml_path(self) -> str:
        project = self.project_path_input.text().strip()
        return os.path.join(project, "output", "dubbing.yaml") if project else ""

    def _dub_audio_path(self, sid: int) -> str:
        project = self.project_path_input.text().strip()
        return os.path.join(project, "output", "audio", f"scene_{sid:03d}.mp3") if project else ""

    def _dub_scenes_yaml_path(self) -> str:
        project = self.project_path_input.text().strip()
        return os.path.join(project, "output", "scenes.yaml") if project else ""

    def _dub_update_total_duration(self) -> None:
        """Show the combined duration of the currently dubbed scene audio."""
        total_seconds = 0.0
        for sid in getattr(self, "_dub_editors", {}):
            audio_path = self._dub_audio_path(sid)
            if not os.path.exists(audio_path):
                continue
            try:
                from mutagen.mp3 import MP3
                total_seconds += MP3(audio_path).info.length
            except Exception:
                try:
                    from moviepy import AudioFileClip
                    clip = AudioFileClip(audio_path)
                    total_seconds += clip.duration
                    clip.close()
                except Exception:
                    pass
        rounded_seconds = max(0, int(total_seconds + 0.5))
        minutes, seconds = divmod(rounded_seconds, 60)
        self._dub_duration_label.setText(f"{minutes:02d}:{seconds:02d}")

    def _dub_segment_speed_text(self, sid: int) -> str:
        """Speed badge text for a scene, based on the TTS rate used for its current audio."""
        if not os.path.exists(self._dub_audio_path(sid)):
            return "—"
        pct = getattr(self, "_dub_rates", {}).get(sid)
        if pct is None:
            return "—"
        return f"{1 + pct / 100:.2f}x"

    def _dub_segment_duration(self, sid: int) -> str:
        """Duration badge text (mm:ss) for a single scene's dubbed audio file."""
        audio_path = self._dub_audio_path(sid)
        if not os.path.exists(audio_path):
            return "--:--"
        seconds = 0.0
        try:
            from mutagen.mp3 import MP3
            seconds = MP3(audio_path).info.length
        except Exception:
            try:
                from moviepy import AudioFileClip
                clip = AudioFileClip(audio_path)
                seconds = clip.duration
                clip.close()
            except Exception:
                return "--:--"
        total = max(0, int(seconds + 0.5))
        minutes, secs = divmod(total, 60)
        return f"{minutes:02d}:{secs:02d}"

    def _dub_update_segment_badges(self, sid: int) -> None:
        """Refresh a single scene card's speed/duration badges."""
        speed_lbl = getattr(self, "_dub_speed_labels", {}).get(sid)
        if speed_lbl:
            speed_lbl.setText(self._dub_segment_speed_text(sid))
        dur_lbl = getattr(self, "_dub_duration_labels", {}).get(sid)
        if dur_lbl:
            dur_lbl.setText(self._dub_segment_duration(sid))

    def _dub_bookmark_path(self) -> str:
        project = self.project_path_input.text().strip()
        return os.path.join(project, "output", "dubbing_bookmark.yaml") if project else ""

    def _dub_state(self, sid: int) -> str:
        """Return 'not_dubbed', 'to_redub', or 'dubbed' for a scene."""
        audio = self._dub_audio_path(sid)
        if not os.path.exists(audio):
            return "not_dubbed"
        if self._dub_dirty.get(sid, False):
            return "to_redub"
        return "dubbed"

    def _dub_apply_card_state(self, sid: int) -> None:
        """Update card border colour and preview button to reflect current state."""
        state = self._dub_state(sid)
        card  = self._dub_cards.get(sid)
        btn   = self._dub_preview_btns.get(sid)
        if not card or not btn:
            return
        bg, col = {
            "not_dubbed": self._DUB_NOT_DUBBED,
            "to_redub":   self._DUB_TO_REDUB,
            "dubbed":     self._DUB_DUBBED,
        }[state]
        border = {
            "not_dubbed": "#36343B",
            "to_redub":   "#F8B23D",
            "dubbed":     "#4CAF50",
        }[state]
        if sid == getattr(self, "_dub_selected_sid", None):
            border = "#96BDE2"
        card.setStyleSheet(
            f"QWidget#dubCard {{ background:#1D1B20; border:2px solid {border}; border-radius:4px; }}"
        )
        icon = {"not_dubbed": "▶", "to_redub": "⟳", "dubbed": "▶"}[state]
        btn.setText(icon)
        btn.setStyleSheet(
            f"QPushButton {{ background:{bg}; color:{col}; border:1px solid {border}; "
            f"border-radius:3px; font-size:14px; "
            f"min-width:28px; max-width:28px; min-height:28px; max-height:28px; padding:0; }}"
            f"QPushButton:hover {{ background:#2A282F; }}"
        )

    def _dub_load_from_scenes(self) -> None:
        scenes_path = self._dub_scenes_yaml_path()
        if not scenes_path or not os.path.exists(scenes_path):
            QMessageBox.warning(self, "No scenes", "Run 'Split Scenes' first.")
            return
        import yaml as _yaml
        scenes = (_yaml.safe_load(Path(scenes_path).read_text(encoding="utf-8")) or {}).get("scenes", [])
        dub_path = self._dub_yaml_path()
        existing: dict = {}
        if dub_path and os.path.exists(dub_path):
            existing = _yaml.safe_load(Path(dub_path).read_text(encoding="utf-8")) or {}
        self._dub_populate(scenes, existing)

    # ── Find / Replace ────────────────────────────────────────────────────────

    def _dub_toggle_find_panel(self) -> None:
        if self._dub_find_panel.isVisible():
            self._dub_close_find_panel()
        else:
            self._dub_open_find_panel()

    def _dub_open_find_panel(self) -> None:
        self._dub_find_panel.setVisible(True)
        self._dub_find_input.setFocus()
        self._dub_find_input.selectAll()
        self._dub_find_cursor = None  # reset position tracker

    def _dub_close_find_panel(self) -> None:
        self._dub_find_panel.setVisible(False)
        self._dub_find_status.setText("")
        self._dub_find_cursor = None

    def _dub_find_editors(self):
        """Return list of (scene_id, QPlainTextEdit) in scene order."""
        if not hasattr(self, "_dub_editors"):
            return []
        return sorted(self._dub_editors.items())

    def _dub_search_flags(self):
        import re
        flags = 0 if self._dub_find_case.isChecked() else re.IGNORECASE
        return flags

    def _dub_find_in_text(self, text: str, query: str) -> list:
        """Return list of (start, end) match spans."""
        import re
        if not query:
            return []
        pattern = re.escape(query)
        if self._dub_find_word.isChecked():
            pattern = r'\b' + pattern + r'\b'
        flags = self._dub_search_flags()
        return [(m.start(), m.end()) for m in re.finditer(pattern, text, flags)]

    def _dub_find_next(self) -> None:
        self._dub_find_step(forward=True)

    def _dub_find_prev(self) -> None:
        self._dub_find_step(forward=False)

    def _dub_find_step(self, forward: bool) -> None:
        query = self._dub_find_input.text()
        if not query:
            return
        editors = self._dub_find_editors()
        if not editors:
            return

        # Build flat list of all matches: (sid, editor, start, end)
        all_matches = []
        for sid, ed in editors:
            for start, end in self._dub_find_in_text(ed.toPlainText(), query):
                all_matches.append((sid, ed, start, end))

        if not all_matches:
            self._dub_find_status.setText("No matches")
            self._dub_find_cursor = None
            return
        self._dub_find_status.setText(f"{len(all_matches)} match(es)")

        # Advance the index
        cur = getattr(self, "_dub_find_cursor", None)
        if cur is None:
            idx = 0 if forward else len(all_matches) - 1
        else:
            if forward:
                idx = (cur + 1) % len(all_matches)
            else:
                idx = (cur - 1) % len(all_matches)

        self._dub_find_cursor = idx
        sid, ed, start, end = all_matches[idx]
        self._dub_highlight_match(ed, start, end)

    def _dub_highlight_match(self, ed, start: int, end: int) -> None:
        from PyQt6.QtGui import QTextCursor
        cursor = ed.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
        ed.setTextCursor(cursor)
        ed.setFocus()
        ed.ensureCursorVisible()
        # Scroll the card into the dubbing scroll area
        card = ed.parentWidget()
        if card and hasattr(self, "_dub_scroll_area"):
            self._dub_scroll_area.ensureWidgetVisible(card)

    def _dub_replace_one(self) -> None:
        from PyQt6.QtGui import QTextCursor
        query = self._dub_find_input.text()
        replacement = self._dub_replace_input.text()
        if not query:
            return
        # Replace current selection if it matches, then find next
        for sid, ed in self._dub_find_editors():
            if ed.hasFocus():
                cursor = ed.textCursor()
                if cursor.hasSelection():
                    selected = cursor.selectedText()
                    import re
                    pattern = re.escape(query)
                    if self._dub_find_word.isChecked():
                        pattern = r'\b' + pattern + r'\b'
                    flags = self._dub_search_flags()
                    if re.fullmatch(pattern, selected, flags):
                        cursor.insertText(replacement)
                        self._dub_dirty[sid] = True
                        self._dub_refresh_status()
                break
        self._dub_find_next()

    def _dub_replace_all(self) -> None:
        import re
        query = self._dub_find_input.text()
        replacement = self._dub_replace_input.text()
        if not query:
            return
        pattern = re.escape(query)
        if self._dub_find_word.isChecked():
            pattern = r'\b' + pattern + r'\b'
        flags = self._dub_search_flags()
        count = 0
        for sid, ed in self._dub_find_editors():
            text = ed.toPlainText()
            new_text, n = re.subn(pattern, replacement, text, flags=flags)
            if n:
                ed.setPlainText(new_text)
                self._dub_dirty[sid] = True
                self._dub_refresh_status()
                count += n
        self._dub_find_status.setText(f"Replaced {count} occurrence(s)")

    def _dub_speed_changed(self, value: float) -> None:
        """Mirror the Dubbing tab speed control into the Settings tab TTS rate field."""
        if not hasattr(self, "tts_rate_input"):
            return
        pct = round((value - 1.0) * 100)
        if self.tts_rate_input.value() != pct:
            self.tts_rate_input.blockSignals(True)
            self.tts_rate_input.setValue(pct)
            self.tts_rate_input.blockSignals(False)

    def _tts_rate_changed(self, value: int) -> None:
        """Mirror the Settings tab TTS rate field into the Dubbing tab speed control."""
        if not hasattr(self, "_dub_speed_input"):
            return
        speed = round(1.0 + value / 100.0, 2)
        if abs(self._dub_speed_input.value() - speed) > 1e-6:
            self._dub_speed_input.blockSignals(True)
            self._dub_speed_input.setValue(speed)
            self._dub_speed_input.blockSignals(False)

    def _dub_export_word(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Open a project first.")
            return
        editors: dict = getattr(self, "_dub_editors", {})
        if not editors:
            QMessageBox.warning(self, "Nothing to export", "Load scenes in the Dubbing tab first.")
            return
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            QMessageBox.critical(self, "Missing dependency",
                                 "python-docx is not installed.\nRun: pip install python-docx")
            return

        doc = Document()
        doc.core_properties.title = os.path.basename(project)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        heading = doc.add_heading(os.path.basename(project), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        for sid in sorted(editors.keys()):
            text = editors[sid].toPlainText().strip()
            doc.add_paragraph(text)

        out_dir = os.path.join(project, "output")
        os.makedirs(out_dir, exist_ok=True)
        project_name = os.path.basename(project).replace(" ", "_")
        out_path = os.path.join(out_dir, f"{project_name}_dubbing.docx")
        try:
            doc.save(out_path)
        except PermissionError:
            QMessageBox.warning(self, "Export failed",
                                f"Could not write:\n{out_path}\n\nClose the file in Word and try again.")
            return
        msg = QMessageBox(self)
        msg.setWindowTitle("Exported")
        msg.setText(f"Word file saved to:\n{out_path}")
        msg.setIcon(QMessageBox.Icon.NoIcon)
        msg.exec()

    def _dub_populate(self, scenes: list, existing: dict = None, from_disk: bool = False) -> None:
        existing = existing or {}
        while self._dub_cards_layout.count() > 1:
            item = self._dub_cards_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._dub_editors: dict       = {}   # sid → QPlainTextEdit
        self._dub_cards:   dict       = {}   # sid → card QWidget
        self._dub_preview_btns: dict  = {}   # sid → QPushButton
        self._dub_dirty:   dict       = {}   # sid → bool (text changed after last dub)
        self._dub_spell_highlighters  = {}   # sid → _SpellHighlighter
        self._dub_speed_labels: dict  = {}   # sid → QLabel (speed badge)
        self._dub_duration_labels: dict = {} # sid → QLabel (duration badge)
        self._dub_rates:   dict       = {}   # sid → TTS rate % used for its current audio
        self._dub_selected_sid = None
        # Load bookmark from disk for this project
        self._dub_bookmarks: dict = {}
        bm_path = self._dub_bookmark_path()
        if bm_path and os.path.exists(bm_path):
            try:
                import yaml as _yaml
                bm_data = _yaml.safe_load(Path(bm_path).read_text(encoding="utf-8")) or {}
                if bm_data.get("bookmarked_scene"):
                    self._dub_bookmarks[int(bm_data["bookmarked_scene"])] = True
            except Exception:
                pass

        for scene in scenes:
            sid      = int(scene["id"])
            original = scene.get("text", "")
            entry    = existing.get(sid, {})
            dubbed   = (entry.get("dubbed") if isinstance(entry, dict) else original) or original
            rate_pct = entry.get("rate_pct") if isinstance(entry, dict) else None
            if rate_pct is not None:
                self._dub_rates[sid] = int(rate_pct)

            card = QWidget()
            card.setObjectName("dubCard")
            cl = QVBoxLayout(card)
            cl.setContentsMargins(10, 8, 10, 8)
            cl.setSpacing(4)

            # Header row: scene label | char count | preview btn
            hdr = QHBoxLayout()
            id_lbl = QLabel(f"Scene {sid}")
            id_lbl.setStyleSheet("color:#96BDE2; font-weight:bold; font-size:11px; background:transparent; border:none;")
            char_lbl = QLabel(f"{len(dubbed)} chars")
            char_lbl.setStyleSheet("color:#8E8B90; font-size:10px; background:transparent; border:none;")
            char_lbl.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

            speed_lbl = QLabel(self._dub_segment_speed_text(sid))
            speed_lbl.setToolTip("Speech speed used to dub this scene")
            speed_lbl.setStyleSheet(
                "color:#C9A6E6; font-size:10px; font-weight:bold; background:#2A2438; "
                "border:1px solid #46365A; border-radius:8px; padding:1px 6px;"
            )

            duration_lbl = QLabel(self._dub_segment_duration(sid))
            duration_lbl.setToolTip("Length of this scene's dubbed audio")
            duration_lbl.setStyleSheet(
                "color:#9FD6B8; font-size:10px; font-weight:bold; background:#1E3A2C; "
                "border:1px solid #2F5B41; border-radius:8px; padding:1px 6px;"
            )

            preview_btn = QPushButton("▶")
            preview_btn.setFixedSize(28, 28)
            preview_btn.setToolTip("Synthesise and play this segment")
            preview_btn.clicked.connect(lambda _, s=sid: self._dub_preview_segment(s))

            deepl_btn = QPushButton("DeepL")
            deepl_btn.setFixedHeight(28)
            deepl_btn.setToolTip("Select this translation and send Ctrl+C twice to DeepL")

            bm_btn = QPushButton("🔖")
            bm_btn.setFixedSize(28, 28)
            bm_btn.setCheckable(True)
            bm_btn.setChecked(self._dub_bookmarks.get(sid, False))
            bm_btn.setToolTip("Toggle bookmark (Ctrl+B cycles through bookmarks)")
            self._dub_apply_bookmark_style(bm_btn, bm_btn.isChecked())
            bm_btn.toggled.connect(lambda checked, s=sid, b=bm_btn: self._dub_toggle_bookmark(s, checked, b))

            hdr.addWidget(id_lbl)
            hdr.addWidget(char_lbl, 1)
            hdr.addWidget(speed_lbl)
            hdr.addWidget(duration_lbl)
            hdr.addWidget(deepl_btn)
            hdr.addWidget(bm_btn)
            hdr.addWidget(preview_btn)
            cl.addLayout(hdr)

            # Original (read-only)
            orig_lbl = QLabel(original)
            orig_lbl.setWordWrap(True)
            orig_lbl.setStyleSheet(
                "color:#5a6470; font-size:12px; background:#131118; "
                "border:1px solid #2a2830; border-radius:2px; padding:4px;"
            )
            cl.addWidget(orig_lbl)

            # Dubbed (editable)
            ed = QPlainTextEdit()
            ed.setPlainText(dubbed)
            ed.setPlaceholderText("Enter dubbed / translated text here…")
            ed.setMinimumHeight(70)
            ed.setMaximumHeight(160)
            ed.setCursorWidth(2)
            ed.setStyleSheet(
                "QPlainTextEdit { background:#1D1B20; color:#E6E1E5; "
                "border:1px solid #36343B; border-radius:2px; padding:4px; "
                "font-family:'Segoe UI',sans-serif; font-size:15px; }"
                "QPlainTextEdit:focus { border:1px solid #96BDE2; }"
            )

            def _on_text_changed(s=sid, lbl=char_lbl, e=ed):
                lbl.setText(f"{len(e.toPlainText())} chars")
                self._dub_dirty[s] = True
                self._dub_has_unsaved = True
                self._dub_apply_card_state(s)
                self._dub_mark_unsaved()

            def _on_focus_in(event, s=sid, e=ed):
                self._dub_select_scene(s)
                QPlainTextEdit.focusInEvent(e, event)

            ed.textChanged.connect(_on_text_changed)
            ed.focusInEvent = _on_focus_in
            ed.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            ed.customContextMenuRequested.connect(
                lambda pos, e=ed: self._dub_show_editor_context_menu(e, pos)
            )
            deepl_btn.clicked.connect(lambda _checked, e=ed: self._dub_send_to_deepl(e))

            cl.addWidget(ed)
            self._dub_editors[sid]      = ed
            self._dub_cards[sid]        = card
            self._dub_preview_btns[sid] = preview_btn
            self._dub_speed_labels[sid]    = speed_lbl
            self._dub_duration_labels[sid] = duration_lbl
            # Attach spell-check highlighter — block signals so rehighlight()
            # does not fire textChanged and spuriously mark the scene as dirty.
            lang = getattr(self, "_dub_spell_lang", "it")
            if lang != "off":
                ed.blockSignals(True)
                hl = _SpellHighlighter(ed.document(), language=lang)
                ed.blockSignals(False)
                self._dub_spell_highlighters[sid] = hl
            self._dub_dirty[sid] = False  # must be set AFTER highlighter creation
            self._dub_cards_layout.insertWidget(self._dub_cards_layout.count() - 1, card)
            self._dub_apply_card_state(sid)

        if self._dub_editors:
            self._dub_select_scene(min(self._dub_editors))

        # Ctrl+B shortcut (created once)
        if not getattr(self, "_dub_bm_shortcut_installed", False):
            from PyQt6.QtGui import QShortcut, QKeySequence
            sc = QShortcut(QKeySequence("Ctrl+B"), self)
            sc.activated.connect(self._dub_goto_next_bookmark)

            self._dub_bm_shortcut_installed = True

        self._dub_status_label.setText(f"{len(scenes)} scene(s) loaded.")
        self._dub_update_total_duration()
        self._dub_has_unsaved = not from_disk
        self._dub_mark_unsaved() if self._dub_has_unsaved else self._dub_mark_saved()

    def _dub_send_to_deepl(self, editor: QPlainTextEdit) -> None:
        """Select one scene translation and invoke DeepL's double-copy shortcut."""
        import ctypes

        editor.setFocus()
        editor.selectAll()

        user32 = ctypes.windll.user32
        key_up = 0x0002
        vk_control = 0x11
        vk_c = 0x43
        for _ in range(2):
            user32.keybd_event(vk_control, 0, 0, 0)
            user32.keybd_event(vk_c, 0, 0, 0)
            user32.keybd_event(vk_c, 0, key_up, 0)
            user32.keybd_event(vk_control, 0, key_up, 0)

    def _dub_mark_unsaved(self) -> None:
        self._dub_update_save_btn()
        self._dub_update_dub_btn()

    def _dub_select_scene(self, sid: int) -> None:
        """Select the scene used as the starting point for sequential playback."""
        previous_sid = getattr(self, "_dub_selected_sid", None)
        self._dub_selected_sid = sid
        self._dub_focused_sid = sid
        if previous_sid in getattr(self, "_dub_cards", {}):
            self._dub_apply_card_state(previous_sid)
        if sid in getattr(self, "_dub_cards", {}):
            self._dub_apply_card_state(sid)
        if hasattr(self, "_dub_status_label"):
            self._dub_status_label.setText(f"Scene {sid} selected for playback.")

    def _dub_mark_saved(self) -> None:
        self._dub_update_save_btn()
        self._dub_update_dub_btn()

    def _dub_apply_bookmark_style(self, btn, active: bool) -> None:
        if active:
            btn.setStyleSheet(
                "QPushButton { background:#4A3800; color:#FFD600; border:1px solid #FFD600; "
                "border-radius:3px; font-size:14px; min-width:28px; min-height:28px; padding:0; }"
                "QPushButton:hover { background:#5A4A00; }"
            )
        else:
            btn.setStyleSheet(
                "QPushButton { background:#2A2830; color:#8E8B90; border:1px solid #36343B; "
                "border-radius:3px; font-size:14px; min-width:28px; min-height:28px; padding:0; }"
                "QPushButton:hover { background:#36343B; }"
            )

    def _dub_toggle_bookmark(self, sid: int, checked: bool, btn) -> None:
        if checked:
            # Remove any existing bookmark from other cards
            for other_sid, active in list(self._dub_bookmarks.items()):
                if active and other_sid != sid:
                    self._dub_bookmarks[other_sid] = False
                    # Update the button visually
                    other_card = self._dub_cards.get(other_sid)
                    if other_card:
                        for child in other_card.findChildren(QPushButton):
                            if child.isCheckable() and child.text() == "🔖":
                                child.blockSignals(True)
                                child.setChecked(False)
                                child.blockSignals(False)
                                self._dub_apply_bookmark_style(child, False)
        self._dub_bookmarks[sid] = checked
        self._dub_apply_bookmark_style(btn, checked)
        # Persist to disk
        bm_path = self._dub_bookmark_path()
        if bm_path:
            import yaml as _yaml
            data = {"bookmarked_scene": sid if checked else None}
            Path(bm_path).write_text(_yaml.dump(data, allow_unicode=True), encoding="utf-8")

    def _dub_goto_next_bookmark(self) -> None:
        bookmarked = [sid for sid, v in getattr(self, "_dub_bookmarks", {}).items() if v]
        if not bookmarked:
            return
        target = bookmarked[0]
        card = getattr(self, "_dub_cards", {}).get(target)
        if not card:
            return
        parent = card.parent()
        while parent:
            from PyQt6.QtWidgets import QScrollArea
            if isinstance(parent, QScrollArea):
                parent.ensureWidgetVisible(card)
                break
            parent = parent.parent()

    def _dub_reload_custom_words(self) -> None:
        """Reload the custom word list from disk and rehighlight all editors."""
        _SpellHighlighter._custom_words = _SpellHighlighter._load_custom_words()
        editors = getattr(self, "_dub_editors", {})
        for sid, hl in getattr(self, "_dub_spell_highlighters", {}).items():
            if hl._checker is not None and _SpellHighlighter._custom_words:
                hl._checker.word_frequency.load_words(_SpellHighlighter._custom_words)
            # Block the editor's signals so rehighlight() cannot trigger
            # _on_text_changed and falsely mark every scene as dirty.
            ed = editors.get(sid)
            if ed:
                ed.blockSignals(True)
            try:
                hl.rehighlight()
            finally:
                if ed:
                    ed.blockSignals(False)
        count = len(_SpellHighlighter._custom_words)
        self._dub_status_label.setText(f"Custom word list reloaded — {count} word(s).")

    def _dub_show_editor_context_menu(self, editor: QPlainTextEdit, pos) -> None:
        """Show the standard editor menu with an Italian dictionary action."""
        cursor = editor.textCursor()
        if not cursor.hasSelection():
            cursor = editor.cursorForPosition(pos)
            cursor.select(QTextCursor.SelectionType.WordUnderCursor)
        word = cursor.selectedText().strip()

        menu = editor.createStandardContextMenu()
        menu.addSeparator()
        add_action = menu.addAction(f'Add "{word}" to Dictionary' if word else "Add to Dictionary")
        import re
        add_action.setEnabled(
            self._dub_spell_lang == "it"
            and bool(re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ']+", word))
            and word.casefold() not in _SpellHighlighter._custom_words
        )
        chosen = menu.exec(editor.viewport().mapToGlobal(pos))
        if chosen == add_action:
            try:
                added = _SpellHighlighter.add_custom_word(word)
            except OSError as exc:
                QMessageBox.warning(self, "Dictionary", f"Could not update the dictionary:\n{exc}")
                return
            self._dub_reload_custom_words()
            if added:
                self._dub_status_label.setText(f'Added "{word}" to the custom dictionary.')

    def _dub_update_spell_btn_style(self) -> None:
        """Color the active IT/EN button; dim the inactive one."""
        active_ss   = ("QPushButton { background:#2A5298; color:#FFFFFF; border:1px solid #4A72B8; "
                       "border-radius:3px; font-size:12px; padding:0 6px; }"
                       "QPushButton:hover { background:#3A63A8; }")
        inactive_ss = ("QPushButton { background:#2A2830; color:#8E8B90; border:1px solid #3A3840; "
                       "border-radius:3px; font-size:12px; padding:0 6px; }"
                       "QPushButton:hover { background:#3A3640; }")
        for key, btn in getattr(self, "_dub_spell_btns", {}).items():
            btn.setStyleSheet(active_ss if key == self._dub_spell_lang else inactive_ss)

    def _dub_set_spell_lang(self, lang: str) -> None:
        self._dub_spell_lang = lang
        self._dub_update_spell_btn_style()
        editors = getattr(self, "_dub_editors", {})
        hls = getattr(self, "_dub_spell_highlighters", {})
        # Update existing or create new highlighters (spell check always on)
        for sid, ed in editors.items():
            ed.blockSignals(True)
            if sid in hls:
                hls[sid].set_language(lang)
            else:
                hls[sid] = _SpellHighlighter(ed.document(), language=lang)
            ed.blockSignals(False)

    def _dub_play_from_current(self) -> None:
        """Play all dubbed audio files sequentially, starting from the focused scene."""
        project = self.project_path_input.text().strip()
        if not project:
            return
        audio_dir = os.path.join(project, "output", "audio")
        editors = getattr(self, "_dub_editors", {})
        if not editors:
            return

        all_sids = sorted(editors.keys())
        start_sid = getattr(self, "_dub_selected_sid", None) or all_sids[0]
        # Collect existing audio files from start_sid onward
        queue = [
            os.path.join(audio_dir, f"scene_{sid:03d}.mp3")
            for sid in all_sids
            if sid >= start_sid and os.path.exists(os.path.join(audio_dir, f"scene_{sid:03d}.mp3"))
        ]
        if not queue:
            self._dub_status_label.setText("No audio files found from this scene onward.")
            return

        self._dub_play_queue = queue
        self._dub_play_index = 0
        self._dub_play_btn.setEnabled(True)
        self._dub_play_btn.setText("⏹ Stop")
        self._dub_play_btn.clicked.disconnect()
        self._dub_play_btn.clicked.connect(self._dub_stop_playback)
        self._dub_advance_playback()

    def _dub_advance_playback(self) -> None:
        queue = getattr(self, "_dub_play_queue", [])
        idx   = getattr(self, "_dub_play_index", 0)
        if idx >= len(queue):
            self._dub_stop_playback()
            return
        from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
        from PyQt6.QtCore import QUrl
        path = queue[idx]
        self._dub_play_index = idx + 1
        # Scroll the corresponding card into view
        try:
            fname = os.path.basename(path)          # scene_003.mp3
            sid = int(fname.split("_")[1].split(".")[0])
            card = self._dub_cards.get(sid)
            if card and hasattr(self, "_dub_scroll_area"):
                self._dub_scroll_area.ensureWidgetVisible(card)
            total = len(queue)
            self._dub_status_label.setText(
                f"Playing scene {sid}  ({idx + 1}/{total})"
            )
        except Exception:
            pass
        if not hasattr(self, "_dub_full_player") or self._dub_full_player is None:
            self._dub_full_player = QMediaPlayer(self)
            self._dub_full_ao = QAudioOutput(self)
            self._dub_full_player.setAudioOutput(self._dub_full_ao)
            self._dub_full_ao.setVolume(1.0)
            self._dub_full_player.playbackStateChanged.connect(self._dub_on_playback_state)
        self._dub_full_player.setSource(QUrl.fromLocalFile(os.path.abspath(path)))
        self._dub_full_player.play()

    def _dub_on_playback_state(self, state) -> None:
        from PyQt6.QtMultimedia import QMediaPlayer
        if state == QMediaPlayer.PlaybackState.StoppedState:
            self._dub_advance_playback()

    def _dub_stop_playback(self) -> None:
        if hasattr(self, "_dub_full_player") and self._dub_full_player:
            # Disconnect BEFORE stop() — stop() emits playbackStateChanged(Stopped)
            # which would otherwise trigger _dub_advance_playback and start the next clip.
            try:
                self._dub_full_player.playbackStateChanged.disconnect()
            except Exception:
                pass
            self._dub_full_player.stop()
            self._dub_full_player = None
        self._dub_play_queue = []
        self._dub_play_btn.setEnabled(True)
        self._dub_play_btn.setText("▶ Play")
        try:
            self._dub_play_btn.clicked.disconnect()
        except Exception:
            pass
        self._dub_play_btn.clicked.connect(self._dub_play_from_current)
        self._dub_status_label.setText("Playback stopped.")

    def _dub_dubbing_status(self) -> str:
        """Compute overall dubbing status for the Dub All button.
        'complete'  — every scene has audio and text is unchanged since last dub
        'stale'     — every scene has audio but at least one text was edited (orange)
        'all_stale' — at least one scene has no audio at all (red)
        'none'      — no scenes loaded
        """
        if not hasattr(self, "_dub_editors") or not self._dub_editors:
            return "none"
        total = len(self._dub_editors)
        has_audio = sum(1 for sid in self._dub_editors if os.path.exists(self._dub_audio_path(sid)))
        if has_audio < total:
            return "all_stale"   # red — missing audio
        # All scenes have audio; check if any text was edited after last dub
        any_dirty = any(self._dub_dirty.get(sid, False) for sid in self._dub_editors)
        return "stale" if any_dirty else "complete"

    def _dub_refresh_status(self) -> None:
        """Refresh dub-all button and save button after programmatic text changes."""
        self._dub_update_dub_btn()
        self._dub_update_save_btn()

    def _dub_update_dub_btn(self) -> None:
        if not hasattr(self, "_dub_all_btn"):
            return
        status = self._dub_dubbing_status()
        bg, hv = self._DUB_BTN_COLORS.get(status, ("#c0392b", "#a93226"))
        self._dub_all_btn.setEnabled(status != "none")
        self._dub_all_btn.setStyleSheet(self._BTN_SS.format(bg=bg, hv=hv))
        if hasattr(self, "_dub_redub_all_btn"):
            self._dub_redub_all_btn.setEnabled(status != "none")

    def _dub_update_save_btn(self) -> None:
        if not hasattr(self, "_dub_save_btn"):
            return
        if not hasattr(self, "_dub_editors") or not self._dub_editors:
            bg, hv, enabled = "#555555", "#444444", False
        elif getattr(self, "_dub_has_unsaved", True):
            bg, hv, enabled = "#c0392b", "#a93226", True   # red — unsaved
        else:
            bg, hv, enabled = "#27ae60", "#1e8449", True   # green — saved
        self._dub_save_btn.setEnabled(enabled)
        self._dub_save_btn.setStyleSheet(self._BTN_SS.format(bg=bg, hv=hv))

    # ── Per-segment preview ───────────────────────────────────────────────────

    def _dub_preview_segment(self, sid: int, auto_next: int = None, silent: bool = False) -> None:
        """Synthesise and play a single segment, save its audio file."""
        self._dub_save()
        ed = self._dub_editors.get(sid)
        if not ed:
            return
        text = ed.toPlainText().strip()
        if not text:
            return

        btn = self._dub_preview_btns.get(sid)
        if btn:
            btn.setEnabled(False)
            btn.setText("…")

        audio_path = self._dub_audio_path(sid)
        os.makedirs(os.path.dirname(audio_path), exist_ok=True)

        voice  = self._tts_voice_setting()
        rate_pct_value = self.tts_rate_input.value()
        rate   = f"{rate_pct_value:+d}%"
        pitch  = f"{self.tts_pitch_input.value():+d}Hz"
        volume = f"{self.tts_volume_input.value():+d}%"

        class _Worker(QObject):
            done  = pyqtSignal(int)
            error = pyqtSignal(int, str)

            def __init__(self, sid, text, voice, rate, pitch, volume, path):
                super().__init__()
                self._sid    = sid
                self._text   = text
                self._voice  = voice
                self._rate   = rate
                self._pitch  = pitch
                self._volume = volume
                self._path   = path

            def run(self):
                try:
                    import sys, subprocess, tempfile as _tf, os as _os
                    kw = ""
                    if self._rate.lstrip('+').rstrip('%') != '0':
                        kw += f"    kwargs['rate'] = {self._rate!r}\n"
                    if self._pitch.lstrip('+').rstrip('Hz') != '0':
                        kw += f"    kwargs['pitch'] = {self._pitch!r}\n"
                    if self._volume.lstrip('+').rstrip('%') != '0':
                        kw += f"    kwargs['volume'] = {self._volume!r}\n"
                    script = (
                        "# -*- coding: utf-8 -*-\nimport asyncio, edge_tts\n"
                        "async def _go():\n    kwargs = {}\n" + kw +
                        f"    c = edge_tts.Communicate({self._text!r}, {self._voice!r}, **kwargs)\n"
                        f"    await c.save({self._path!r})\nasyncio.run(_go())\n"
                    )
                    sf = _tf.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8')
                    sf.write(script); sf.close()
                    try:
                        r = subprocess.run([sys.executable, sf.name],
                                           capture_output=True, text=True, timeout=60)
                    finally:
                        _os.unlink(sf.name)
                    if r.returncode != 0:
                        raise RuntimeError(r.stderr.strip() or "Synthesis failed")
                    self.done.emit(self._sid)
                except Exception as exc:
                    self.error.emit(self._sid, str(exc))

        thread = QThread(self)
        worker = _Worker(sid, text, voice, rate, pitch, volume, audio_path)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)

        def _on_done(s):
            self._dub_dirty[s] = False
            self._dub_rates[s] = rate_pct_value
            self._dub_apply_card_state(s)
            self._dub_update_dub_btn()
            self._dub_update_total_duration()
            self._dub_update_segment_badges(s)
            self._dub_save()
            b = self._dub_preview_btns.get(s)
            if b:
                b.setEnabled(True)
            if not silent:
                from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
                from PyQt6.QtCore import QUrl
                self._seg_player = QMediaPlayer(self)
                ao = QAudioOutput(self); self._seg_player.setAudioOutput(ao)
                self._seg_audio_out = ao; ao.setVolume(1.0)
                self._seg_player.setSource(QUrl.fromLocalFile(os.path.abspath(audio_path)))
                self._seg_player.play()
            if auto_next is not None:
                self._dub_all_next(auto_next)
            thread.quit()

        def _on_error(s, msg):
            b = self._dub_preview_btns.get(s)
            if b:
                b.setEnabled(True)
                self._dub_apply_card_state(s)
            if auto_next is not None:
                self._dub_all_next(auto_next)   # continue even on error
            else:
                QMessageBox.warning(self, f"Scene {s} synthesis failed", msg)
            thread.quit()

        worker.done.connect(_on_done)
        worker.error.connect(_on_error)
        if not hasattr(self, "_dub_seg_threads"):
            self._dub_seg_threads = []
        self._dub_seg_threads.append((thread, worker))
        thread.start()

    def _tts_voice_setting(self) -> str:
        try:
            return self.tts_voice_input.currentData() or "it-IT-DiegoNeural"
        except Exception:
            return "it-IT-DiegoNeural"

    def _on_tts_voice_changed(self) -> None:
        """Mark all segments that have existing audio as dirty when voice changes."""
        if not hasattr(self, "_dub_editors") or not self._dub_editors:
            return
        changed = False
        for sid in self._dub_editors:
            if os.path.exists(self._dub_audio_path(sid)):
                self._dub_dirty[sid] = True
                self._dub_apply_card_state(sid)
                changed = True
        if changed:
            self._dub_has_unsaved = True
            self._dub_update_dub_btn()
            self._dub_update_save_btn()
            self._dub_status_label.setText("Voice changed — re-dub segments to update audio.")

    # ── Dub All ───────────────────────────────────────────────────────────────

    def _dub_all(self, force: bool = False) -> None:
        if not hasattr(self, "_dub_editors") or not self._dub_editors:
            QMessageBox.warning(self, "No segments", "Load scenes first.")
            return
        if force:
            # Redub every segment regardless of dirty/existing-audio state
            # (e.g. after changing speed, pitch or voice for the whole project).
            ids = sorted(self._dub_editors.keys())
        else:
            # Only queue segments that have no audio yet or whose text has changed
            ids = sorted(
                sid for sid in self._dub_editors
                if not os.path.exists(self._dub_audio_path(sid)) or self._dub_dirty.get(sid, False)
            )
        if not ids:
            self._dub_status_label.setText("All segments already dubbed and up to date.")
            return
        self._dub_all_queue = ids
        self._dub_all_total = len(ids)
        # Track progress relative to ALL segments so the bar shows meaningful
        # movement even when only a few segments remain.
        self._dub_all_grand_total = len(self._dub_editors)
        self._dub_all_already_done = self._dub_all_grand_total - self._dub_all_total
        self._dub_all_btn.setEnabled(False)
        self._dub_redub_all_btn.setEnabled(False)
        self._dub_progress.setVisible(True)
        start_pct = int(self._dub_all_already_done / self._dub_all_grand_total * 100)
        self._dub_progress.setValue(start_pct)
        from PyQt6.QtWidgets import QApplication
        QApplication.processEvents()
        self._dub_all_next(-1)

    def _dub_all_next(self, last_done_index: int) -> None:
        done = last_done_index + 1
        total = self._dub_all_total
        grand_total = getattr(self, "_dub_all_grand_total", total)
        already_done = getattr(self, "_dub_all_already_done", 0)
        if done >= total:
            self._dub_all_btn.setEnabled(True)
            self._dub_redub_all_btn.setEnabled(True)
            self._dub_progress.setVisible(False)
            self._dub_generate_timings()
            self._dub_save()
            self._dub_update_dub_btn()
            self._dub_status_label.setText(f"Dub All complete — {total} segment(s) synthesised.")
            return
        pct = int((already_done + done) / grand_total * 100)
        self._dub_progress.setValue(pct)
        sid = self._dub_all_queue[done]
        self._dub_preview_segment(sid, auto_next=done, silent=True)

    # ── Save ──────────────────────────────────────────────────────────────────

    def _dub_generate_timings(self) -> None:
        """Measure duration of each dubbed MP3 and write timings.yaml."""
        project = self.project_path_input.text().strip()
        if not project:
            return
        audio_dir = os.path.join(project, "output", "audio")
        timings_path = os.path.join(audio_dir, "timings.yaml")
        timings: dict = {}
        for sid in sorted(self._dub_editors.keys() if hasattr(self, "_dub_editors") else []):
            mp3 = os.path.join(audio_dir, f"scene_{sid:03d}.mp3")
            if not os.path.exists(mp3):
                continue
            try:
                from mutagen.mp3 import MP3
                timings[sid] = MP3(mp3).info.length
            except Exception:
                try:
                    from moviepy import AudioFileClip
                    c = AudioFileClip(mp3); dur = c.duration; c.close()
                    timings[sid] = dur
                except Exception:
                    pass
        if timings:
            import yaml as _yaml
            os.makedirs(audio_dir, exist_ok=True)
            Path(timings_path).write_text(
                _yaml.dump(timings, sort_keys=True), encoding="utf-8"
            )
            self._append_log(f"timings.yaml written: {len(timings)} scene(s).")

    def _dub_save(self) -> None:
        path = self._dub_yaml_path()
        if not path:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        if not hasattr(self, "_dub_editors") or not self._dub_editors:
            QMessageBox.warning(self, "Nothing to save", "Load scenes first.")
            return
        scenes_path = self._dub_scenes_yaml_path()
        originals: dict = {}
        if scenes_path and os.path.exists(scenes_path):
            import yaml as _yaml
            scenes = (_yaml.safe_load(Path(scenes_path).read_text(encoding="utf-8")) or {}).get("scenes", [])
            originals = {int(s["id"]): s.get("text", "") for s in scenes}
        rates = getattr(self, "_dub_rates", {})
        data = {
            sid: {
                "original": originals.get(sid, ""),
                "dubbed": ed.toPlainText(),
                **({"rate_pct": rates[sid]} if sid in rates else {}),
            }
            for sid, ed in self._dub_editors.items()
        }
        try:
            import yaml as _yaml
            os.makedirs(os.path.dirname(path), exist_ok=True)
            Path(path).write_text(
                _yaml.dump(data, allow_unicode=True, sort_keys=True, default_flow_style=False),
                encoding="utf-8",
            )
            self._dub_status_label.setText(f"Saved {len(data)} scene(s) — {path}")
            self._dub_has_unsaved = False
            self._dub_mark_saved()
        except Exception as exc:
            QMessageBox.critical(self, "Save failed", str(exc))

    def _dub_refresh(self) -> None:
        dub_path = self._dub_yaml_path()
        scenes_path = self._dub_scenes_yaml_path()
        if not scenes_path or not os.path.exists(scenes_path):
            self._dub_update_save_btn()
            self._dub_update_dub_btn()
            return
        import yaml as _yaml
        scenes = (_yaml.safe_load(Path(scenes_path).read_text(encoding="utf-8")) or {}).get("scenes", [])
        existing: dict = {}
        if dub_path and os.path.exists(dub_path):
            existing = _yaml.safe_load(Path(dub_path).read_text(encoding="utf-8")) or {}
        self._dub_populate(scenes, existing, from_disk=True)

    def _build_script_tab(self) -> QWidget:
        page = QWidget()
        root = QVBoxLayout(page)
        root.setSpacing(6)

        header_row = QHBoxLayout()
        self._script_path_label = QLabel("No project loaded")
        self._script_path_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        header_row.addWidget(self._script_path_label, 1)

        header_row.addWidget(QLabel("Text size:"))
        self._script_font_size = 15
        self._script_font_size_slider = QSlider(Qt.Orientation.Horizontal)
        self._script_font_size_slider.setMinimum(8)
        self._script_font_size_slider.setMaximum(36)
        self._script_font_size_slider.setValue(self._script_font_size)
        self._script_font_size_slider.setFixedWidth(120)
        self._script_font_size_slider.setToolTip("Adjust script text size")
        self._script_font_size_slider.valueChanged.connect(self._on_script_font_size_changed)
        header_row.addWidget(self._script_font_size_slider)
        self._script_font_size_label = QLabel(f"{self._script_font_size}px")
        self._script_font_size_label.setStyleSheet("color:#8E8B90; font-size:11px;")
        self._script_font_size_label.setFixedWidth(28)
        header_row.addWidget(self._script_font_size_label)
        root.addLayout(header_row)

        self._script_editor = QPlainTextEdit()
        self._script_editor.setPlaceholderText("Open or create a project to edit the narration script.")
        self._script_editor.textChanged.connect(self._on_script_text_changed)
        root.addWidget(self._script_editor, 1)
        self._apply_script_font_size(self._script_font_size)

        btn_row = QHBoxLayout()
        self._btn_save_script = QPushButton("Save Script")
        self._btn_save_script.clicked.connect(self._save_script)
        btn_reload_script = QPushButton("Reload from Disk")
        btn_reload_script.clicked.connect(self._reload_script)
        self._btn_sync_script = QPushButton("⟳ Sync to Dubbing")
        self._btn_sync_script.setToolTip(
            "Save script → re-split scenes → generate missing prompts (Ollama) → refresh Dubbing tab.\n"
            "Images are NOT generated — run Preview Images (stage 5) to check prompts first."
        )
        self._btn_sync_script.clicked.connect(self._sync_script_to_pipeline)
        btn_row.addStretch(1)
        btn_row.addWidget(btn_reload_script)
        btn_row.addWidget(self._btn_save_script)
        btn_row.addWidget(self._btn_sync_script)
        root.addLayout(btn_row)

        self._script_dirty = False
        self._script_set_saved_style()
        return page

    def _apply_script_font_size(self, size: int) -> None:
        self._script_editor.setStyleSheet(
            f"QPlainTextEdit {{ font-family: 'Segoe UI', sans-serif; font-size: {size}px; line-height: 1.5; }}"
        )

    def _on_script_font_size_changed(self, size: int) -> None:
        self._script_font_size = size
        self._script_font_size_label.setText(f"{size}px")
        self._apply_script_font_size(size)

    def _script_set_saved_style(self) -> None:
        self._btn_save_script.setStyleSheet(
            "QPushButton { background:#1e4620; color:#4CAF50; font-weight:bold; "
            "border:1px solid #4CAF50; border-radius:4px; padding:4px 14px; }"
            "QPushButton:hover { background:#27ae60; color:white; }"
        )
        self._btn_save_script.setText("Save Script")

    def _script_set_dirty_style(self) -> None:
        self._btn_save_script.setStyleSheet(
            "QPushButton { background:#4a1010; color:#e74c3c; font-weight:bold; "
            "border:1px solid #e74c3c; border-radius:4px; padding:4px 14px; }"
            "QPushButton:hover { background:#c0392b; color:white; }"
        )
        self._btn_save_script.setText("Save Script ●")

    def _on_script_text_changed(self) -> None:
        if not self._script_dirty:
            self._script_dirty = True
            self._script_set_dirty_style()

    def _script_file_path(self) -> str:
        project = self.project_path_input.text().strip()
        if not project:
            return ""
        return os.path.join(project, "input", "narration.txt")

    def _reload_script(self) -> None:
        path = self._script_file_path()
        if not path:
            return
        if os.path.exists(path):
            # Block textChanged so reload doesn't trigger dirty flag
            self._script_editor.blockSignals(True)
            self._script_editor.setPlainText(Path(path).read_text(encoding="utf-8"))
            self._script_editor.blockSignals(False)
            self._script_path_label.setText(path)
        else:
            self._script_editor.blockSignals(True)
            self._script_editor.clear()
            self._script_editor.blockSignals(False)
            self._script_path_label.setText(f"{path}  (not found)")
        self._script_dirty = False
        self._script_set_saved_style()

    def _save_script(self) -> None:
        path = self._script_file_path()
        if not path:
            QMessageBox.warning(self, "No project", "Load a project before saving the script.")
            return
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            Path(path).write_text(self._script_editor.toPlainText(), encoding="utf-8")
            self._script_path_label.setText(f"Saved: {path}")
            self._script_dirty = False
            self._script_set_saved_style()
        except Exception as exc:
            QMessageBox.critical(self, "Save failed", str(exc))

    def _sync_script_to_pipeline(self) -> None:
        """Save → split scenes → refresh Dubbing → run missing prompts+images."""
        # 1. Save first
        self._save_script()
        if self._script_dirty:
            return  # save failed

        # 2. Re-split scenes via pipeline (runs synchronously is not possible;
        #    run it as a pipeline stage so the user sees progress)
        reply = QMessageBox.question(
            self, "Sync to Dubbing",
            "This will:\n"
            "  1. Re-split scenes from the saved script\n"
            "  2. Generate missing prompts via Ollama\n"
            "  3. Refresh the Dubbing tab\n\n"
            "Existing scenes and prompts are preserved.\n"
            "Images are NOT generated — run Preview Images (stage 5) to check prompts.\n\n"
            "Continue?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return

        # Switch to Pipeline tab so the user can see progress
        for i in range(self.tabs.count()):
            if self.tabs.tabText(i) == "Pipeline":
                self.tabs.setCurrentIndex(i)
                break

        self._sync_pending = True
        self.controller.run_pipeline("prompts")  # narration → scenes → prompts only (no images)

    def _wire_signals(self) -> None:
        self.controller.project_changed.connect(self._on_project_changed)
        self.controller.recent_projects_changed.connect(self._refresh_recent_projects)
        self.controller.pipeline_started.connect(self._on_pipeline_started)
        self.controller.pipeline_progress.connect(self._on_pipeline_progress)
        self.controller.pipeline_log.connect(self._append_log)
        self.controller.pipeline_finished.connect(self._on_pipeline_finished)
        self.controller.settings_saved.connect(self._on_settings_saved)

    def _on_project_changed(self, path: str) -> None:
        # Only reload script from disk if the editor is empty or project actually changed
        current_project = self.project_path_input.text().strip()
        project_changed = os.path.abspath(path) != os.path.abspath(current_project) if current_project else True
        self.project_path_input.setText(path)
        if project_changed:
            self._reload_script()
        self._refresh_draft_grid()
        self._refresh_lightbox()
        self._dub_refresh()
        self._prompts_refresh()

    def _refresh_recent_projects(self, projects) -> None:
        self.recent_projects_combo.blockSignals(True)
        self.recent_projects_combo.clear()
        self.recent_projects_combo.addItems(projects)
        self.recent_projects_combo.blockSignals(False)

    def _on_pipeline_started(self) -> None:
        for btn in getattr(self, "_stage_btns", {}).values():
            btn.setEnabled(False)
        self.btn_cancel_pipeline.setEnabled(True)
        self.pipeline_status_label.setText("Pipeline started")
        self.pipeline_progress.setValue(0)
        self._append_log("Pipeline started.")

    def _on_pipeline_progress(self, value: int, message: str) -> None:
        self.pipeline_progress.setValue(value)
        self.pipeline_status_label.setText(message)
        self._append_log(message)

    def _append_log(self, message: str) -> None:
        self.log_output.appendPlainText(message)

    @staticmethod
    def _flush_cuda_main_thread() -> None:
        """
        Called from the main thread (via QTimer) after a GPU pipeline stage
        completes.  Forces Python GC and frees the CUDA allocator cache so
        that VRAM drops back to idle.  NOTE: torch.cuda.synchronize() is
        intentionally NOT called here — calling it from the Qt event-loop
        thread (a different thread from the one that submitted the kernels)
        can hard-crash the CUDA driver on Windows, making the process
        disappear silently.
        """
        import gc
        gc.collect()
        try:
            import torch
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                gc.collect()
        except Exception:
            pass

    def _on_pipeline_finished(self, success: bool, payload: str) -> None:
        for btn in getattr(self, "_stage_btns", {}).values():
            btn.setEnabled(True)
        self.btn_cancel_pipeline.setEnabled(False)

        # Flush any leftover CUDA work from the worker thread on the main
        # thread so the GPU drops back to idle without waiting for a window
        # focus event to trigger the event-loop flush.
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(0, self._flush_cuda_main_thread)

        # Auto-refresh the preview grid whenever preview_images or final_images completes.
        if success and payload and os.path.isdir(payload) and (
            "draft" in payload or "images" in payload
        ):
            # Defer so we never rebuild the grid inside a nested event-loop
            # (e.g. while a zoom dialog is still open).
            QTimer.singleShot(0, self._refresh_draft_grid)

        # Auto-refresh the Lightbox tab when final_images completes.
        if success and payload and os.path.isdir(payload) and "lightbox" in payload:
            QTimer.singleShot(0, self._refresh_lightbox)

        # After a sync-triggered prompts run, refresh the Dubbing tab
        if success and hasattr(self, "_sync_pending") and self._sync_pending:
            self._sync_pending = False
            self._dub_refresh()

        if success:
            QTimer.singleShot(0, self._prompts_refresh)
            if payload and os.path.basename(payload) == "model_prompts.yaml":
                QTimer.singleShot(0, self._refresh_draft_grid)
                QTimer.singleShot(0, self._refresh_lightbox)

        if success:
            self.pipeline_status_label.setText("Pipeline complete")
            self.pipeline_progress.setValue(100)
            self._append_log(f"Finished: {payload}")
            return

        if "canceled" in payload.lower():
            self.pipeline_status_label.setText("Pipeline canceled")
            self._append_log(payload)
            QMessageBox.information(self, "Pipeline canceled", payload)
            return

        self.pipeline_status_label.setText("Pipeline failed")
        self._append_log(f"Error: {payload}")
        QMessageBox.critical(self, "Pipeline failed", payload)

    def _on_settings_saved(self, path: str) -> None:
        QMessageBox.information(self, "Settings saved", f"Settings saved to:\n{path}")

    def _load_settings_to_form(self, settings: dict) -> None:
        self.style_preset_input.setCurrentText(str(settings.get("style_preset", "cinematic")))
        self.aspect_ratio_input.setCurrentText(str(settings.get("aspect_ratio", "16:9")))
        self.seed_input.setValue(int(settings.get("seed", 42)))
        self.fps_input.setValue(int(settings.get("fps", 8)))
        self.scene_split_method_input.setCurrentText(str(settings.get("scene_split_method", "paragraph")))
        self.min_sentence_length_input.setValue(int(settings.get("min_sentence_length", 20)))
        self.sdxl_model_input.setText(str(settings.get("sdxl_base", "models/sd3")))
        self.flux_dev_model_input.setText(str(settings.get("flux_dev", "models/flux/FLUX.1-dev")))
        self.flux_schnell_model_input.setText(str(settings.get("flux_schnell", "models/flux/FLUX.1-schnell")))
        self.flux2_model_input.setText(str(settings.get("flux2", "models/flux/FLUX.2-klein-4B")))
        self.svd_model_input.setText(str(settings.get("svd", "models/svd")))
        self.clip_engine_input.setCurrentText(str(settings.get("clip_engine", "ken_burns")))
        self.ken_burns_motion_input.setCurrentText(str(settings.get("ken_burns_motion", "static")))
        self.schnell_steps_input.setValue(int(settings.get("schnell_steps", 4)))
        self.schnell_guidance_input.setValue(float(settings.get("schnell_guidance", 0.0)))
        self.dev_steps_input.setValue(int(settings.get("dev_steps", 20)))
        self.dev_guidance_input.setValue(float(settings.get("dev_guidance", 3.5)))
        self.flux2_steps_input.setValue(int(settings.get("flux2_steps", 4)))
        self.flux2_guidance_input.setValue(float(settings.get("flux2_guidance", 1.0)))
        w = int(settings.get("image_width", 1024))
        h = int(settings.get("image_height", 576))
        idx = self.image_resolution_input.findData((w, h))
        self.image_resolution_input.setCurrentIndex(idx if idx >= 0 else 0)
        self.num_frames_input.setValue(int(settings.get("num_frames", 14)))
        self.motion_bucket_id_input.setValue(int(settings.get("motion_bucket_id", 127)))
        self.audio_volume_input.setValue(float(settings.get("audio_volume", 1.0)))
        self.fade_in_input.setValue(float(settings.get("fade_in", 0.0)))
        self.fade_out_input.setValue(float(settings.get("fade_out", 0.0)))
        self.use_ollama_input.setChecked(bool(settings.get("use_ollama", False)))
        self.ollama_model_input.setText(str(settings.get("ollama_model", "qwen3:8b")))
        self.ollama_host_input.setText(str(settings.get("ollama_host", "http://localhost:11434")))

        # TTS
        tts_voice = str(settings.get("tts_voice", "it-IT-DiegoNeural"))
        idx = self.tts_voice_input.findData(tts_voice)
        if idx >= 0:
            self.tts_voice_input.setCurrentIndex(idx)
        tts_rate = str(settings.get("tts_rate", "+0%")).replace("+", "").replace("%", "")
        self.tts_rate_input.setValue(int(tts_rate) if tts_rate.lstrip("-").isdigit() else 0)
        self._tts_rate_changed(self.tts_rate_input.value())
        tts_pitch = str(settings.get("tts_pitch", "+0Hz")).replace("+", "").replace("Hz", "")
        self.tts_pitch_input.setValue(int(tts_pitch) if tts_pitch.lstrip("-").isdigit() else 0)
        tts_volume = str(settings.get("tts_volume", "+0%")).replace("+", "").replace("%", "")
        self.tts_volume_input.setValue(int(tts_volume) if tts_volume.lstrip("-").isdigit() else 0)

    def _collect_settings_from_form(self) -> dict:
        return {
            "style_preset": self.style_preset_input.currentText(),
            "aspect_ratio": self.aspect_ratio_input.currentText(),
            "seed": self.seed_input.value(),
            "fps": self.fps_input.value(),
            "scene_split_method": self.scene_split_method_input.currentText(),
            "min_sentence_length": self.min_sentence_length_input.value(),
            "sdxl_base": self.sdxl_model_input.text().strip(),
            "flux_dev": self.flux_dev_model_input.text().strip(),
            "flux_schnell": self.flux_schnell_model_input.text().strip(),
            "flux2": self.flux2_model_input.text().strip(),
            "svd": self.svd_model_input.text().strip(),
            "clip_engine": self.clip_engine_input.currentText(),
            "ken_burns_motion": self.ken_burns_motion_input.currentText(),
            "schnell_steps": self.schnell_steps_input.value(),
            "schnell_guidance": self.schnell_guidance_input.value(),
            "dev_steps": self.dev_steps_input.value(),
            "dev_guidance": self.dev_guidance_input.value(),
            "flux2_steps": self.flux2_steps_input.value(),
            "flux2_guidance": self.flux2_guidance_input.value(),
            "guidance_scale": self.schnell_guidance_input.value(),  # compat
            "num_inference_steps": self.schnell_steps_input.value(),  # compat
            "image_model": "flux-schnell",  # always schnell for preview now
            "image_width":  self.image_resolution_input.currentData()[0],
            "image_height": self.image_resolution_input.currentData()[1],
            "num_frames": self.num_frames_input.value(),
            "motion_bucket_id": self.motion_bucket_id_input.value(),
            "audio_volume": self.audio_volume_input.value(),
            "fade_in": self.fade_in_input.value(),
            "fade_out": self.fade_out_input.value(),
            "use_ollama": self.use_ollama_input.isChecked(),
            "ollama_model": self.ollama_model_input.text().strip() or "qwen3:8b",
            "ollama_host": self.ollama_host_input.text().strip() or "http://localhost:11434",
            "tts_voice": self.tts_voice_input.currentData() or "it-IT-DiegoNeural",
            "tts_rate": f"{self.tts_rate_input.value():+d}%",
            "tts_pitch": f"{self.tts_pitch_input.value():+d}Hz",
            "tts_volume": f"{self.tts_volume_input.value():+d}%",
            # pass-through: no dedicated UI widget, preserve current config value
            "output_width": int(self.controller.config.get("output_width", 1920)),
            "output_height": int(self.controller.config.get("output_height", 1080)),
        }

    def select_project(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Project Folder")
        if folder:
            self.controller.set_project_path(folder)

    def use_recent_project(self):
        selected = self.recent_projects_combo.currentText().strip()
        if selected:
            self.controller.set_project_path(selected)

    def create_project(self):
        parent = QFileDialog.getExistingDirectory(self, "Select Parent Folder")
        if not parent:
            return

        name, accepted = QInputDialog.getText(self, "Create Project", "Project name:")
        if not accepted or not name.strip():
            return

        try:
            created = self.controller.create_project(parent, name.strip())
            QMessageBox.information(self, "Project created", f"Created project:\n{created}")
        except Exception as exc:
            QMessageBox.critical(self, "Create project failed", str(exc))

    def _preview_tts_voice(self) -> None:
        voice  = self.tts_voice_input.currentData() or "it-IT-DiegoNeural"
        rate   = f"{self.tts_rate_input.value():+d}%"
        pitch  = f"{self.tts_pitch_input.value():+d}Hz"
        volume = f"{self.tts_volume_input.value():+d}%"

        self._tts_preview_btn.setEnabled(False)
        self._tts_preview_btn.setText("…")

        import tempfile, asyncio

        tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
        tmp_path = tmp.name
        tmp.close()

        class _Worker(QObject):
            done  = pyqtSignal(str)
            error = pyqtSignal(str)

            def __init__(self, voice, rate, pitch, volume, path):
                super().__init__()
                self._voice  = voice
                self._rate   = rate
                self._pitch  = pitch
                self._volume = volume
                self._path   = path

            def run(self):
                try:
                    import sys, subprocess, tempfile as _tf, os as _os
                    kwargs_lines = ""
                    if self._rate.lstrip('+').rstrip('%') != '0':
                        kwargs_lines += f"    kwargs['rate'] = {self._rate!r}\n"
                    if self._pitch.lstrip('+').rstrip('Hz') != '0':
                        kwargs_lines += f"    kwargs['pitch'] = {self._pitch!r}\n"
                    if self._volume.lstrip('+').rstrip('%') != '0':
                        kwargs_lines += f"    kwargs['volume'] = {self._volume!r}\n"
                    script = (
                        "# -*- coding: utf-8 -*-\n"
                        "import asyncio, edge_tts\n"
                        "async def _go():\n"
                        "    kwargs = {}\n"
                        + kwargs_lines +
                        f"    c = edge_tts.Communicate('Ciao, questa e una anteprima della voce.', {self._voice!r}, **kwargs)\n"
                        f"    await c.save({self._path!r})\n"
                        "asyncio.run(_go())\n"
                    )
                    sf = _tf.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8')
                    sf.write(script)
                    sf.close()
                    try:
                        result = subprocess.run(
                            [sys.executable, sf.name],
                            capture_output=True, text=True, timeout=30,
                        )
                    finally:
                        _os.unlink(sf.name)
                    if result.returncode != 0:
                        raise RuntimeError(result.stderr.strip() or "Synthesis failed")
                    self.done.emit(self._path)
                except Exception as exc:
                    self.error.emit(str(exc))

        self._tts_preview_thread = QThread(self)
        self._tts_preview_worker = _Worker(voice, rate, pitch, volume, tmp_path)
        self._tts_preview_worker.moveToThread(self._tts_preview_thread)
        self._tts_preview_thread.started.connect(self._tts_preview_worker.run)

        def _on_done(path):
            self._tts_preview_btn.setEnabled(True)
            self._tts_preview_btn.setText("▶ Preview")
            from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
            from PyQt6.QtCore import QUrl
            self._tts_preview_player = QMediaPlayer(self)
            audio_out = QAudioOutput(self)
            self._tts_preview_player.setAudioOutput(audio_out)
            self._tts_preview_audio_out = audio_out
            audio_out.setVolume(1.0)
            self._tts_preview_player.setSource(QUrl.fromLocalFile(path))
            self._tts_preview_player.play()
            self._tts_preview_thread.quit()

        def _on_error(msg):
            self._tts_preview_btn.setEnabled(True)
            self._tts_preview_btn.setText("▶ Preview")
            QMessageBox.warning(self, "TTS Preview failed", msg)
            self._tts_preview_thread.quit()

        self._tts_preview_worker.done.connect(_on_done)
        self._tts_preview_worker.error.connect(_on_error)
        self._tts_preview_thread.start()

    def _relaunch(self) -> None:
        if os.environ.get("VID_COMFY_LAUNCHER") == "1":
            QCoreApplication.exit(75)
            return

        subprocess.Popen([sys.executable] + sys.argv, cwd=os.getcwd())
        QCoreApplication.quit()

    def _clear_prompt_cache(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        path = os.path.join(project, "output", "prompts.yaml")
        if not os.path.exists(path):
            QMessageBox.information(self, "No cache", "No prompts.yaml found — nothing to clear.")
            return
        try:
            os.remove(path)
            self._draft_status_label.setText("Prompt cache cleared — next run will regenerate all prompts.")
        except Exception as exc:
            QMessageBox.critical(self, "Delete failed", str(exc))

    def run_stage(self, stage: str):
        path = self.project_path_input.text().strip()
        if path:
            self.controller.set_project_path(path)
        self.controller.run_pipeline(stage)

    def run_selected_stage(self):
        """Legacy — kept for any remaining references."""
        pass

    def _clear_preview_clips(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        clips_dir = os.path.join(project, "output", "draft_clips")
        if not os.path.isdir(clips_dir):
            QMessageBox.information(self, "No clips", "No draft_clips folder found.")
            return
        files = [f for f in os.listdir(clips_dir) if f.lower().endswith(".mp4")]
        if not files:
            QMessageBox.information(self, "No clips", "draft_clips folder is already empty.")
            return
        reply = QMessageBox.question(
            self, "Clear Preview Clips",
            f"Delete {len(files)} preview clip(s)?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        errors = []
        for f in files:
            try:
                os.remove(os.path.join(clips_dir, f))
            except Exception as exc:
                errors.append(str(exc))
        if errors:
            QMessageBox.warning(self, "Some files not deleted", "\n".join(errors))
        else:
            self._append_log(f"Cleared {len(files)} preview clip(s) from output/draft_clips/.")

    def _clear_lightbox(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        lightbox_dir = os.path.join(project, "output", "lightbox")
        if not os.path.isdir(lightbox_dir):
            QMessageBox.information(self, "No lightbox", "No lightbox folder found.")
            return
        files = [f for f in os.listdir(lightbox_dir)
                 if f.lower().endswith((".png", ".jpg", ".jpeg"))]
        if not files:
            QMessageBox.information(self, "No images", "Lightbox folder is already empty.")
            return
        reply = QMessageBox.question(
            self, "Clear Lightbox",
            f"Delete {len(files)} lightbox image(s)?\nThey will need to be regenerated with step 8.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        errors = []
        for f in files:
            try:
                os.remove(os.path.join(lightbox_dir, f))
            except Exception as exc:
                errors.append(str(exc))
        if errors:
            QMessageBox.warning(self, "Some files not deleted", "\n".join(errors))
        else:
            self._append_log(f"Cleared {len(files)} lightbox image(s) from output/lightbox/.")
            self._refresh_lightbox()

    def _clear_draft(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        draft_dir = os.path.join(project, "output", "draft")
        if not os.path.isdir(draft_dir):
            QMessageBox.information(self, "No draft", "No draft folder found.")
            return
        files = [f for f in os.listdir(draft_dir) if f.lower().endswith((".png", ".jpg", ".jpeg"))]
        if not files:
            QMessageBox.information(self, "No draft", "Draft folder is already empty.")
            return
        reply = QMessageBox.question(
            self, "Clear Draft Images",
            f"Delete {len(files)} draft image(s)?\nThey will be regenerated using the current prompts.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        for f in files:
            try:
                os.remove(os.path.join(draft_dir, f))
            except Exception:
                pass
        self._refresh_draft_grid()
        self._append_log(f"Cleared {len(files)} draft image(s).")

    def _clear_clips(self) -> None:
        project = self.project_path_input.text().strip()
        if not project:
            QMessageBox.warning(self, "No project", "Load a project first.")
            return
        clips_dir = os.path.join(project, "output", "clips")
        if not os.path.isdir(clips_dir):
            QMessageBox.information(self, "No clips", "No clips folder found.")
            return
        files = [f for f in os.listdir(clips_dir) if f.lower().endswith(".mp4")]
        if not files:
            QMessageBox.information(self, "No clips", "Clips folder is already empty.")
            return
        reply = QMessageBox.question(
            self, "Clear Clips",
            f"Delete {len(files)} clip(s) from output/clips/?\nThey will be regenerated on the next run.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        errors = []
        for f in files:
            try:
                os.remove(os.path.join(clips_dir, f))
            except Exception as exc:
                errors.append(str(exc))
        if errors:
            QMessageBox.warning(self, "Some files not deleted", "\n".join(errors))
        else:
            self._append_log(f"Cleared {len(files)} clip(s) from output/clips/.")

    def cancel_pipeline(self):
        self.controller.cancel_pipeline()

    def save_settings(self):
        try:
            self.controller.save_settings(self._collect_settings_from_form())
        except Exception as exc:
            import traceback, sys
            traceback.print_exc(file=sys.stderr)
            QMessageBox.critical(self, "Save settings failed", traceback.format_exc())
